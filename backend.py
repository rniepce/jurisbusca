import os
import re
import traceback
import tempfile
import hashlib
import time
import random
import concurrent.futures
from typing import List, Optional, Any
import pypdf
import docx
from langchain_community.document_loaders import PyPDFLoader

# --- Imports Condicionais (podem falhar no Railway sem certas dependências) ---

# OCR Engine
HAS_OCR = False
try:
    import ocr_engine
    HAS_OCR = True
    print("✅ ocr_engine importado com sucesso.")
except ImportError as e:
    print(f"⚠️ ocr_engine não disponível (ImportError): {e}")
except Exception as e:
    print(f"⚠️ Erro ao importar ocr_engine: {e}")
    traceback.print_exc()

# Hybrid Chunker
HybridSemanticChunker = None
try:
    from chunking import HybridSemanticChunker
except ImportError:
    print("⚠️ HybridSemanticChunker não disponível.")
except Exception as e:
    print(f"⚠️ Erro ao importar HybridSemanticChunker: {e}")

# RAPTOR Engine
RaptorEngine = None
try:
    from raptor_engine import RaptorEngine
except ImportError:
    print("⚠️ RaptorEngine não disponível.")
except Exception as e:
    print(f"⚠️ Erro ao importar RaptorEngine: {e}")

# Planning Engine
PlanningEngine = None
try:
    from planning_engine import PlanningEngine
except ImportError:
    print("⚠️ PlanningEngine não disponível.")
except Exception as e:
    print(f"⚠️ Erro ao importar PlanningEngine: {e}")

# Style Engine
StyleEngine = None
try:
    from style_engine import StyleEngine
except ImportError:
    print("⚠️ StyleEngine não disponível.")
except Exception as e:
    print(f"⚠️ Erro ao importar StyleEngine: {e}")

# Agent Workflow
create_agent_workflow = None
try:
    from agent_workflow import create_agent_workflow
except ImportError:
    print("⚠️ create_agent_workflow não disponível.")
except Exception as e:
    print(f"⚠️ Erro ao importar create_agent_workflow: {e}")


from langchain_community.vectorstores import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.messages import SystemMessage, HumanMessage, AIMessage
import json
import gc

# Provider Integrations
try:
    import google.generativeai as genai
    from langchain_google_genai import ChatGoogleGenerativeAI
    HAS_GEMINI = True
    GEMINI_IMPORT_ERROR = None
except ImportError as e:
    HAS_GEMINI = False
    GEMINI_IMPORT_ERROR = str(e)
    ChatGoogleGenerativeAI = None

try:
    from langchain_openai import ChatOpenAI, OpenAIEmbeddings
    HAS_OPENAI = True
except ImportError:
    HAS_OPENAI = False
    ChatOpenAI = None
    OpenAIEmbeddings = None

try:
    from langchain_openai import AzureChatOpenAI
    HAS_AZURE_OPENAI = True
except ImportError:
    HAS_AZURE_OPENAI = False
    AzureChatOpenAI = None

# Legacy flag kept for compatibility checks elsewhere
HAS_ANTHROPIC = HAS_AZURE_OPENAI
from prompts import PROMPT_FATOS, PROMPT_ANALISE_FORMAL, PROMPT_JUIZ_DEEPSEEK, PROMPT_REDATOR_CLAUDE, PROMPT_AUDITOR_GPT
from prompts_claude import PROMPT_CLAUDE_INTEGRAL, PROMPT_GPT_AUDITOR, PROMPT_STYLE_ANALYZER, PROMPT_XRAY_BATCH, PROMPT_GPT_FIXER

# V2 Imports (Agentic)
try:
    from v2_engine.orchestrator_v2 import run_hybrid_orchestration
    from v3_engine.orchestrator_v3 import run_autonomous_magistrate
except ImportError as e:
    # Se falhar (ex: falta langgraph), apenas V2 ficará indisponível
    print(f"Erro ao importar V2/V3 Engine: {e}")
    run_hybrid_orchestration = None
    run_autonomous_magistrate = None

# ── Module-level cache for style dossier (avoids re-running expensive analysis) ──
# Keyed by (user_id, file_hash) for per-user isolation
_style_dossier_cache = {}

def _template_cache_key(template_files, user_id="default"):
    """Gera chave de cache baseada nos nomes dos arquivos de template + user_id."""
    names = sorted([getattr(f, 'name', str(f)) for f in template_files])
    return hashlib.md5('|'.join(names).encode()).hexdigest()


def safe_content(response) -> str:
    """
    Normaliza response.content para string limpa.
    Lida com TODOS os formatos de retorno:
    - Anthropic: lista [{'type':'text','text':'...'}]
    - Gemini 2.5: lista [{'type':'thinking','thinking':'...'},{'type':'text','text':'...'}]
    - OpenAI: string pura
    - LangChain ChatGeneration: .content pode ser string ou lista
    - Casos edge: None, dict, int, etc.
    """
    # Extrair .content de objetos LangChain (AIMessage, ChatGeneration, etc.)
    if hasattr(response, 'content'):
        content = response.content
    elif hasattr(response, 'text'):
        content = response.text
    elif isinstance(response, dict):
        content = response.get('text', response.get('content', response.get('output', '')))
    else:
        content = response

    # None → string vazia
    if content is None:
        return ""

    # Se já é string, limpa e retorna
    if isinstance(content, str):
        content = content.replace("\\n", "\n")
        return content.strip()

    # Se é lista (Anthropic, Gemini 2.5 com thinking blocks)
    if isinstance(content, list):
        parts = []
        for item in content:
            if isinstance(item, dict):
                # Pula blocos de "thinking" do modelo (raciocínio interno)
                if item.get('type') == 'thinking':
                    continue
                if 'text' in item:
                    parts.append(str(item['text']))
                elif 'content' in item:
                    parts.append(str(item['content']))
                else:
                    parts.append(str(item))
            elif isinstance(item, str):
                parts.append(item)
            elif hasattr(item, 'text'):
                parts.append(str(item.text))
            elif hasattr(item, 'content'):
                parts.append(str(item.content))
            else:
                parts.append(str(item))
        content = "\n".join(parts)
        content = content.replace("\\n", "\n")
        return content.strip()

    # Se é dict (ex: JSON response acidental)
    if isinstance(content, dict):
        if 'text' in content:
            return str(content['text']).strip()
        if 'content' in content:
            return str(content['content']).strip()
        # Serializa como JSON legível
        import json as _json
        return _json.dumps(content, ensure_ascii=False, indent=2)

    # Fallback total: converte para string
    result = str(content)
    result = result.replace("\\n", "\n")
    return result.strip()


def clean_llm_text(text: str) -> str:
    """
    Sanitiza texto vindo de LLM para exibição limpa.
    Remove tags HTML, escaped newlines, artefatos de dict/JSON.
    """
    if not text or not isinstance(text, str):
        return ""
    # 1. Converte escaped newlines
    text = text.replace("\\n", "\n")
    # 2. Remove tags HTML comuns que LLMs podem gerar
    text = re.sub(r'<(?!/?(?:br|hr)\s*/?>)[^>]+>', '', text)
    # 3. Remove artefatos de dict Python vazando
    if "'extras':" in text:
        text = text.split("'extras':")[0].strip().rstrip(",").strip()
    elif '"extras":' in text:
        text = text.split('"extras":')[0].strip().rstrip(",").strip()
    # 4. Remove aspas soltas no início/fim
    text = text.strip().strip("'").strip('"')
    return text


def clean_text(text: str) -> str:
    """
    Higienização agressiva para peças jurídicas (Otimização de Context Window).
    Remove: Cabeçalhos, Rodapés, Números de Página, Espaços duplos, Assinaturas Digitais.
    """
    if not text or not isinstance(text, str):
        return ""
    
    original_text = text
    original_len = len(text)
        
    # 1. Normalização de quebras de linha
    text = text.replace('\r', '')
    
    # 2. (Preservado) Números de processo e folhas são juridicamente essenciais — não remover

    # 3. Remove rodapés de escritório/sistema e assinaturas digitais
    # Padrão comum: "PJe - Assinado eletronicamente" ou "Documento assinado digitalmente"
    text = re.sub(r'(?i)(assinado\s+eletronicamente|documento\s+assinado|pje|assinatura\s+digital).*', '', text) 
    
    # 4. Remove números de página soltos
    text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
    
    # 5. Redução de ruído visual (traços, asteriscos)
    text = re.sub(r'[_=\-\*]{3,}', '', text)
    
    # 6. NOVO: Remove blocos de assinatura digital Base64 (longas sequências alfanuméricas)
    # Detecta strings com mais de 200 caracteres consecutivos sem espaços (típico de Base64/hash)
    text = re.sub(r'[A-Za-z0-9+/=]{200,}', '', text)
    
    # 7. NOVO: Remove chaves/colchetes JSON com conteúdo de 'signature' ou 'extras'
    text = re.sub(r"'extras':\s*\{[^}]*\}", '', text)
    text = re.sub(r"'signature':\s*'[^']*'", '', text)
    
    # 8. NOVO: Remove linhas que parecem ser metadados de certificado
    text = re.sub(r'(?i)(certificado|hash|sha\d*|md5|rsa|dsa|asn\.\d):[^\n]*\n?', '', text)
    
    # 9. Compressão de espaços (White space normalization) - PRESERVANDO quebras de linha
    # Substitui múltiplos espaços horizontais por um único
    text = re.sub(r'[ \t]+', ' ', text)
    # Limita múltiplas quebras de linha a no máximo duas
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    cleaned = text.strip()
    
    # Safety check: if cleaning removed more than 80% of content,
    # fall back to minimal cleaning to avoid stripping substantive text
    if original_len > 100 and len(cleaned) < original_len * 0.2:
        # Minimal cleaning: only whitespace normalization
        minimal = re.sub(r'[ \t]+', ' ', original_text)
        minimal = re.sub(r'\n{3,}', '\n\n', minimal)
        return minimal.strip()
    
    return cleaned

def get_embedding_function(api_key=None):
    """
    Factory centralizada de embeddings — usa Azure OpenAI (text-embedding-3-small).
    api_key: se fornecida, tem prioridade sobre a variável de ambiente.
    """
    from langchain_openai import AzureOpenAIEmbeddings
    
    azure_key = api_key or os.getenv("AZURE_OPENAI_API_KEY", "")
    azure_endpoint = os.getenv("AZURE_OPENAI_EMBEDDING_ENDPOINT", os.getenv("AZURE_OPENAI_ENDPOINT", ""))
    deployment = os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT", "text-embedding-3-small")
    
    if not azure_key or not azure_endpoint:
        raise ValueError(
            "AZURE_OPENAI_API_KEY e AZURE_OPENAI_EMBEDDING_ENDPOINT devem estar configurados "
            "no .env para usar embeddings."
        )
    
    return AzureOpenAIEmbeddings(
        azure_deployment=deployment,
        azure_endpoint=azure_endpoint,
        api_key=azure_key,
        api_version="2024-12-01-preview",
    )

def process_uploaded_file(file_obj, filename: str, api_key=None, ocr_engine_choice="mistral_doc_ai", compress=True, vectorize=True, progress_callback=None):
    """
    Salva arquivo temp, faz OCR se necessário, vetoriza e retorna (full_text, retriever).
    progress_callback(msg, percent) is called at each stage for UI progress bar.
    """
    def _progress(msg, pct):
        if progress_callback:
            progress_callback(msg, pct)
        print(msg)

    text = ""
    docs = []
    
    # Cria arquivo temporário para processamento (necessário para loaders do Langchain)
    suffix = os.path.splitext(filename)[1].lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
        tmp_file.write(file_obj.read())
        tmp_path = tmp_file.name

    try:
        if suffix == ".pdf":
            if ocr_engine_choice == "none":
                # ── No OCR: direct text extraction only (PyPDFLoader) ──
                _progress("📄 Extraindo texto nativo do PDF...", 30)
                from langchain_community.document_loaders import PyPDFLoader
                loader = PyPDFLoader(tmp_path)
                docs = loader.load()
                total_chars = sum(len(d.page_content) for d in docs)
                _progress(f"📄 Texto nativo extraído: {len(docs)} páginas, {total_chars:,} caracteres", 50)
            else:
                # ── Hybrid Extract: page-level triage (text vs OCR) ──
                try:
                    from core.hybrid_extract import hybrid_extract
                    ocr_choice = ocr_engine_choice
                    # Normalize: any invalid engine defaults to mistral_doc_ai
                    if ocr_choice not in ("marker", "mistral_doc_ai", "tesseract"):
                        ocr_choice = "mistral_doc_ai"
                    _progress(f"🔍 Analisando páginas com OCR ({ocr_choice})...", 30)
                    docs, stats = hybrid_extract(tmp_path, ocr_choice, compress)
                    _progress(
                        f"📊 OCR concluído: {stats['text_pages']} págs texto + {stats['ocr_pages']} págs OCR ({stats['total_chars']:,} chars)",
                        55
                    )
                except ImportError:
                    _progress("⚠️ hybrid_extract não disponível. Usando extração legada...", 30)
                    # ── Fallback: extração legada (PyPDFLoader) ──
                    from langchain_community.document_loaders import PyPDFLoader
                    loader = PyPDFLoader(tmp_path)
                    docs = loader.load()
                    total_chars = sum(len(d.page_content) for d in docs)
                    if total_chars < 500:
                        _progress(f"📉 Texto insuficiente ({total_chars} chars). Acionando OCR...", 40)
                        if HAS_OCR:
                            try:
                                from ocr_engine import get_marker_engine
                                marker = get_marker_engine()
                                if marker:
                                    ocr_text = marker.process_pdf(tmp_path)
                                    if ocr_text:
                                        from langchain_core.documents import Document
                                        docs = [Document(page_content=ocr_text, metadata={"source": filename, "ocr": "marker"})]
                            except Exception as ocr_err:
                                print(f"⚠️ OCR Marker fallback falhou: {ocr_err}")
                    _progress(f"📄 Extração concluída: {len(docs)} documentos", 50)
        
        elif suffix == ".docx":
            _progress("📝 Extraindo texto do DOCX...", 30)
            from langchain_community.document_loaders import Docx2txtLoader
            loader = Docx2txtLoader(tmp_path)
            docs = loader.load()
            _progress(f"📝 DOCX extraído: {len(docs)} seções", 50)
            
        elif suffix == ".txt":
            _progress("📝 Lendo arquivo de texto...", 30)
            from langchain_community.document_loaders import TextLoader
            try:
                loader = TextLoader(tmp_path, encoding='utf-8')
                docs = loader.load()
            except Exception:
                loader = TextLoader(tmp_path, encoding='latin-1')
                docs = loader.load()
            _progress(f"📝 TXT lido: {len(docs)} bloco(s)", 50)
            
        else:
            raise ValueError(f"Formato não suportado: {suffix} ({filename})")

        # Limpeza e Consolidação
        _progress("🧹 Limpando e consolidando texto...", 55)
        full_text = ""
        for doc in docs:
            cleaned = clean_text(doc.page_content)
            doc.page_content = cleaned
            full_text += cleaned + "\n\n"
            
        _progress(f"📏 Texto extraído: {len(full_text):,} caracteres", 60)

        # Skip vectorization if not requested
        if not vectorize:
            _progress("⏭️ Vetorização desativada. Retornando texto sem RAG.", 95)
            return full_text, None

        # Vetorização (RAG)
        # Divide em chunks semânticos (seções jurídicas + embeddings)
        if not docs:
             return full_text, None

        try:
            from chunking import HybridSemanticChunker
            _progress("🧠 Criando chunks semânticos (seções jurídicas)...", 65)
            semantic_chunker = HybridSemanticChunker(api_key=api_key)
            splits = semantic_chunker.split_text(
                full_text,
                source_metadata={"source": filename}
            )
            if splits:
                # Log section distribution
                sections = {}
                for s in splits:
                    sec = s.metadata.get("section", "GERAL")
                    sections[sec] = sections.get(sec, 0) + 1
                _progress(f"🧠 {len(splits)} chunks semânticos gerados ({len(sections)} seções)", 75)
            else:
                raise ValueError("Chunker retornou lista vazia")
        except Exception as chunk_err:
            _progress(f"⚠️ Chunking semântico falhou. Usando fallback por caracteres...", 70)
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=4000,
                chunk_overlap=200,
                separators=["\n\n", "\n", " ", ""]
            )
            splits = text_splitter.split_documents(docs)
            _progress(f"📦 {len(splits)} chunks por caracteres criados", 75)
        
        # Cria Vector Store em memória (ephemeral)
        try:
            _progress(f"🔗 Vetorizando {len(splits)} chunks com embeddings...", 80)
            embedding_function = get_embedding_function(api_key=api_key)

            # Timeout para a vetorização inteira (5 min max) — usa thread, compatível com background tasks
            RAG_TIMEOUT = 300

            def _do_vectorize():
                vs = Chroma.from_documents(
                    documents=splits,
                    embedding=embedding_function,
                    collection_name="temp_process_analysis"
                )
                return vs.as_retriever(search_kwargs={"k": 15})

            try:
                with concurrent.futures.ThreadPoolExecutor(max_workers=1) as _executor:
                    _future = _executor.submit(_do_vectorize)
                    retriever = _future.result(timeout=RAG_TIMEOUT)
            except concurrent.futures.TimeoutError:
                raise TimeoutError(f"Vetorização excedeu {RAG_TIMEOUT}s")
            
            _progress(f"✅ RAG indexado: {len(splits)} chunks vetorizados", 95)
            return full_text, retriever
            
        except TimeoutError as e:
            _progress(f"⚠️ RAG Timeout. Retornando texto sem vetorização.", 95)
            return full_text, None
        except Exception as e:
            _progress(f"⚠️ Erro na vetorização. Retornando texto sem retriever.", 95)
            return full_text, None
        
    except Exception as e:
        traceback.print_exc()
        raise RuntimeError(f"Erro ao processar arquivo '{filename}': {type(e).__name__}: {e}") from e
    finally:
        # Limpa arquivo temporário
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

def get_llm(model_name: str = "gpt-5.3-chat", temperature: float = 0.2, api_key: str = None, **kwargs):
    """
    Factory centralizada — suporta Azure OpenAI, Google Gemini e Anthropic Claude.
    model_name: 'gpt-5.3-chat' (Azure), 'gemini-3.1-pro' (Google), 'claude-sonnet-4-6' (Anthropic).
    """
    deployment = model_name or os.getenv("AZURE_OPENAI_DEPLOYMENT", "gpt-5.3-chat")

    # ── Google Gemini (native API) ──
    if deployment.startswith("gemini"):
        try:
            from langchain_google_genai import ChatGoogleGenerativeAI
        except ImportError:
            raise ImportError("langchain-google-genai não instalado. Execute: pip install langchain-google-genai")
        
        google_key = os.getenv("GOOGLE_API_KEY", "")
        if not google_key:
            raise ValueError("GOOGLE_API_KEY deve estar configurada para usar Gemini.")
        
        # Map deployment names to Google model IDs
        gemini_model_map = {
            "gemini-3.1-pro": "gemini-3.1-pro-preview",
            "gemini-2.5-pro": "gemini-2.5-pro",
            "gemini-2.5-flash": "gemini-2.5-flash",
        }
        google_model = gemini_model_map.get(deployment, deployment)
        
        # Gemini 2.5+ and 3.x support thinking — enable dynamic thinking
        gemini_thinking_models = {"gemini-2.5-pro", "gemini-2.5-flash", "gemini-3.1-pro-preview"}
        if google_model in gemini_thinking_models:
            if 'thinking_budget' not in kwargs:
                kwargs['thinking_budget'] = -1  # Dynamic thinking (model decides depth)
            if 'max_output_tokens' not in kwargs:
                kwargs['max_output_tokens'] = 16384
            print(f"🟢🧠 Gemini Thinking: {google_model} | thinking_budget={kwargs['thinking_budget']} | max_output={kwargs['max_output_tokens']}")
        else:
            print(f"🟢 Gemini: {google_model}")
        
        return ChatGoogleGenerativeAI(
            model=google_model,
            google_api_key=google_key,
            temperature=temperature,
            **kwargs,
        )

    # ── Anthropic Claude (native API) ──
    if deployment.startswith("claude"):
        try:
            from langchain_anthropic import ChatAnthropic
        except ImportError:
            raise ImportError("langchain-anthropic não instalado. Execute: pip install langchain-anthropic")
        
        anthropic_key = os.getenv("ANTHROPIC_API_KEY", "")
        if not anthropic_key:
            raise ValueError("ANTHROPIC_API_KEY deve estar configurada para usar Claude.")
        
        # Map deployment names to Anthropic model IDs
        claude_model_map = {
            "claude-sonnet-4-6": "claude-4-6-sonnet-20260220",
            "claude-sonnet-4-5": "claude-sonnet-4-5-20250514",
        }
        anthropic_model = claude_model_map.get(deployment, deployment)
        
        # Enable extended thinking for Claude Sonnet 4+ models
        # Extended thinking requires temperature=1 (Anthropic constraint)
        claude_thinking_models = {"claude-4-6-sonnet-20260220", "claude-sonnet-4-5-20250514"}
        if anthropic_model in claude_thinking_models:
            if 'thinking' not in kwargs:
                kwargs['thinking'] = {"type": "enabled", "budget_tokens": 10000}
            print(f"🟠🧠 Claude Extended Thinking: {anthropic_model} | budget={kwargs['thinking'].get('budget_tokens', 'N/A')}")
            return ChatAnthropic(
                model=anthropic_model,
                anthropic_api_key=anthropic_key,
                temperature=1,  # Required by Anthropic when thinking is enabled
                max_tokens=16384,
                **kwargs,
            )
        
        print(f"🟠 Claude: {anthropic_model}")
        return ChatAnthropic(
            model=anthropic_model,
            anthropic_api_key=anthropic_key,
            temperature=temperature,
            max_tokens=16384,
            **kwargs,
        )

    # ── Standard Azure OpenAI models (GPT-5.3, etc.) ──
    if not HAS_AZURE_OPENAI:
        raise ImportError("langchain-openai não instalado. Execute: pip install langchain-openai")

    azure_key = api_key or os.getenv("AZURE_OPENAI_API_KEY", "")
    azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT", "")

    if not azure_key or not azure_endpoint:
        raise ValueError(
            "AZURE_OPENAI_API_KEY e AZURE_OPENAI_ENDPOINT devem estar configurados no .env ou variáveis de ambiente."
        )

    # Azure OpenAI uses max_completion_tokens instead of max_tokens
    if 'max_tokens' in kwargs:
        kwargs['max_completion_tokens'] = kwargs.pop('max_tokens')

    # GPT-5.3 is a reasoning model — configure reasoning_effort and output budget
    reasoning_models = {"gpt-5.3-chat"}
    if deployment in reasoning_models:
        # Enable full reasoning power
        if 'reasoning_effort' not in kwargs:
            kwargs['reasoning_effort'] = 'high'
        # Default max_completion_tokens to 16384 (reasoning tokens + output tokens)
        if 'max_completion_tokens' not in kwargs:
            kwargs['max_completion_tokens'] = 16384
        print(f"🧠 Reasoning model: {deployment} | effort={kwargs['reasoning_effort']} | max_tokens={kwargs['max_completion_tokens']}")

    # GPT-5.3 doesn't support custom temperature — only default (1)
    models_no_temp = {"gpt-5.3-chat"}
    use_temperature = temperature if deployment not in models_no_temp else None

    llm_kwargs = dict(
        azure_deployment=deployment,
        azure_endpoint=azure_endpoint,
        api_key=azure_key,
        api_version="2024-12-01-preview",
        **kwargs,
    )
    if use_temperature is not None:
        llm_kwargs["temperature"] = use_temperature

    return AzureChatOpenAI(**llm_kwargs)

def run_reflexion_loop(draft_text, source_text, api_key):
    """
    ACTIVE AUDITOR (REFLEXION LOOP):
    1. Auditor: Critica a minuta (busca alucinações).
    2. Fixer: Se houver erro, reescreve a minuta e devolve.
    """
    try:
        # Usa Claude Sonnet via Azure para auditoria
        auditor_llm = get_llm("gpt-5.3-chat", temperature=0.0)
        
        # 1. Auditoria
        # Precisamos parsear o draft. Se for JSON, extraímos a 'minuta_final'.
        # Se for string (fallback), usamos ela mesma.
        draft_content = draft_text
        if isinstance(draft_text, str) and draft_text.strip().startswith("{"):
            try:
                import json
                # Tenta limpar wrappers
                clean = draft_text.replace("```json", "").replace("```", "").strip()
                data = json.loads(clean)
                if isinstance(data, dict):
                    draft_content = data.get("minuta_final", draft_text)
            except Exception:
                pass

        print("🛡️ Iniciando Auditoria Ativa (Reflexion Loop)...")
        msg_audit = [
            SystemMessage(content=PROMPT_GPT_AUDITOR),
            HumanMessage(content=f"DADOS DO PROCESSO:\n{source_text[:50000]}\n\nMINUTA PARA AUDITORIA:\n{draft_content}")
        ]
        
        audit_resp = safe_content(auditor_llm.invoke(msg_audit))
        audit_clean = audit_resp.replace("```json", "").replace("```", "").strip()
        
        audit_json = {}
        try:
            audit_json = json.loads(audit_clean)
        except Exception:
            print(f"Erro parse auditoria: {audit_clean}")
            return draft_text, "Falha no Parse da Auditoria"

        # 2. Decisão: Aprova ou Corrige?
        if audit_json.get("aprovado") is True:
            print("✅ Auditoria Aprovada (Sem Alucinações).")
            return draft_text, audit_resp # Retorna original
            
        else:
            errors = audit_json.get("erros_criticos", [])
            print(f"❌ Auditoria Reprovou. Erros: {errors}. Iniciando Auto-Correção...")
            
            # 3. Fixer (Usa o mesmo modelo)
            fixer_llm = get_llm("gpt-5.3-chat", temperature=0.1)
            
            msg_fix = PROMPT_GPT_FIXER.format(
                draft=draft_content,
                critique=json.dumps(errors, ensure_ascii=False)
            )
            
            fix_resp = safe_content(fixer_llm.invoke([HumanMessage(content=msg_fix)]))
            
            # Se o input era JSON, precisamos reconstruir o JSON com a minuta corrigida?
            # Sim, para mater compatibilidade com o frontend que espera JSON.
            if isinstance(draft_text, str) and draft_text.strip().startswith("{"):
                try:
                    clean = draft_text.replace("```json", "").replace("```", "").strip()
                    data = json.loads(clean)
                    if isinstance(data, dict):
                        data["minuta_final"] = fix_resp
                        data["diagnostico"]["status_auditoria"] = "Corrigido Automaticamente"
                        return json.dumps(data, ensure_ascii=False), audit_resp
                except Exception:
                    pass
            
            return fix_resp, audit_resp

    except Exception as e:
        print(f"Erro no Reflexion Loop: {e}")
        return draft_text, str(e)




def _extract_template_texts(template_files):
    """
    Lightweight text extraction from template files (PDF/DOCX/TXT).
    Returns a list of Document objects with full text — NO chunking, NO ChromaDB.
    Used by generate_style_dossier which only needs raw text for LLM analysis.
    """
    from langchain_core.documents import Document
    documents = []
    
    for file in template_files:
        if hasattr(file, 'seek'):
            file.seek(0)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file.name.split('.')[-1]}") as tmp:
            tmp.write(file.read())
            tmp_path = tmp.name
        
        try:
            text = ""
            if file.name.endswith(".pdf"):
                reader = pypdf.PdfReader(tmp_path)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            elif file.name.endswith(".docx"):
                doc = docx.Document(tmp_path)
                text = "\n".join([p.text for p in doc.paragraphs])
            else:  # txt
                try:
                    with open(tmp_path, "r", encoding="utf-8") as f:
                        text = f.read()
                except UnicodeDecodeError:
                    with open(tmp_path, "r", encoding="latin-1") as f:
                        text = f.read()
            
            if text.strip():
                documents.append(Document(
                    page_content=text,
                    metadata={"source": file.name}
                ))
        finally:
            os.remove(tmp_path)
    
    return documents

def generate_style_dossier(template_files, api_key):
    """
    ETAPA 1 DO PIPELINE FORENSE: Gera o Dossiê de Identidade Decisional.
    
    Executa a análise dos 5 Pilares uma única vez (cacheado por set de templates).
    Retorna dict com: 'dossier', 'glossary', 'cloning_prompt', 'full_response'.
    """
    # Check cache first
    cache_key = _template_cache_key(template_files)
    if cache_key in _style_dossier_cache:
        print(f"✅ Dossiê de estilo recuperado do cache (key: {cache_key[:8]}...)")
        return _style_dossier_cache[cache_key]
    # Also support per-user cache lookup
    user_cache_key = f"user:{cache_key}"  # Will be overridden by caller if needed
    
    if not HAS_AZURE_OPENAI:
        print("⚠️ Azure OpenAI não instalado. Dossiê de estilo indisponível.")
        return {"error": "Azure OpenAI não instalado. Dossiê de estilo não disponível."}
    
    try:
        print("🧬 Gerando Dossiê de Identidade Decisional (5 Pilares)...")
        
        # Reset seek position on all template files (BytesIO safety)
        for f in template_files:
            if hasattr(f, 'seek'):
                f.seek(0)
        
        # 1. Extrair texto dos templates (sem ChromaDB/RAG — desnecessário para dossiê)
        all_docs = _extract_template_texts(template_files)
        
        if not all_docs:
            print("⚠️ Nenhum documento extraído dos templates.")
            return {"error": "Nenhum texto pôde ser extraído dos arquivos enviados. Verifique se são PDFs com texto (não imagens escaneadas)."}
        
        # 2. Concatenar TODOS os textos (não random sampling como antes)
        # Limita a ~80k chars para caber no context window do Flash
        full_text = ""
        for doc in all_docs:
            source = doc.metadata.get('source', 'desconhecido')
            full_text += f"\n\n=== DECISÃO ({source}) ===\n{doc.page_content}\n"
            if len(full_text) > 80000:
                full_text += "\n[... documentos adicionais truncados por limite de contexto ...]\n"
                break
        
        # 3. Chamar LLM com o prompt forense de 5 pilares
        llm = get_llm("gpt-5.3-chat", temperature=0.3, max_tokens=8000)
        
        messages = [
            SystemMessage(content=PROMPT_STYLE_ANALYZER),
            HumanMessage(content=f"Aqui está o acervo de decisões do magistrado para análise forense:\n{full_text}")
        ]
        
        response = llm.invoke(messages)
        content = safe_content(response)
        
        # 4. Parse structured response into 3 parts
        result = _parse_dossier_response(content)
        result['full_response'] = content
        
        # 5. Cache result
        _style_dossier_cache[cache_key] = result
        
        print(f"✅ Dossiê gerado com sucesso:")
        print(f"   - Dossiê: {len(result.get('dossier', ''))} chars")
        print(f"   - Glossário: {len(result.get('glossary', ''))} chars")
        print(f"   - System Prompt: {len(result.get('cloning_prompt', ''))} chars")
        
        return result
        
    except Exception as e:
        print(f"❌ Erro ao gerar Dossiê de Estilo: {e}")
        traceback.print_exc()
        return {"error": f"Erro ao gerar dossiê: {str(e)}"}

def _parse_dossier_response(content: str) -> dict:
    """
    Parseia a resposta estruturada do prompt forense usando os delimitadores.
    Retorna dict com 'dossier', 'glossary', 'cloning_prompt'.
    """
    result = {'dossier': '', 'glossary': '', 'cloning_prompt': ''}
    
    # Try structured parsing with delimiters
    if '===PARTE_1_DOSSIE===' in content:
        parts = content.split('===PARTE_1_DOSSIE===')
        if len(parts) > 1:
            rest = parts[1]
            
            if '===PARTE_2_GLOSSARIO===' in rest:
                dossier_part, rest2 = rest.split('===PARTE_2_GLOSSARIO===', 1)
                result['dossier'] = dossier_part.strip()
                
                if '===PARTE_3_SYSTEM_PROMPT===' in rest2:
                    glossary_part, rest3 = rest2.split('===PARTE_3_SYSTEM_PROMPT===', 1)
                    result['glossary'] = glossary_part.strip()
                    
                    # Remove trailing ===FIM=== if present
                    cloning = rest3.split('===FIM===')[0] if '===FIM===' in rest3 else rest3
                    result['cloning_prompt'] = cloning.strip()
                else:
                    result['glossary'] = rest2.strip()
            else:
                result['dossier'] = rest.strip()
    else:
        # Fallback: treat entire response as dossier (unstructured)
        print("⚠️ Resposta não contém delimitadores estruturados. Usando como dossiê completo.")
        result['dossier'] = content
        result['cloning_prompt'] = content  # Use full content as cloning instruction
    
    return result

def retrieve_mirror_context(text, api_key, template_files, style_dossier=None):
    """
    Estratégia do Espelho APRIMORADA.
    Combina: System Prompt de Clonagem (do dossiê) + Golden Sample (RAG top-1).
    
    Args:
        text: Texto do processo sendo analisado
        api_key: Chave de API para embeddings
        template_files: Arquivos de modelo de decisão
        style_dossier: Dict com 'dossier', 'glossary', 'cloning_prompt' (opcional)
    """
    if not template_files: return ""
    
    try:
        rag_context = ""
        
        # ── ETAPA 1: Injetar System Prompt de Clonagem (estilo geral) ──
        if style_dossier:
            cloning_prompt = style_dossier.get('cloning_prompt', '')
            glossary = style_dossier.get('glossary', '')
            
            if cloning_prompt:
                rag_context += "\n\n## 🧬 SYSTEM PROMPT DE CLONAGEM (ESTILO DO MAGISTRADO)\n"
                rag_context += "⚠️ INSTRUÇÃO PRIMÁRIA: Você DEVE seguir rigorosamente as diretrizes abaixo para replicar o estilo deste magistrado.\n"
                rag_context += f"\n{cloning_prompt}\n"
            
            if glossary:
                rag_context += "\n\n## 📝 GLOSSÁRIO DO MAGISTRADO (CHECKLIST DE VOCABULÁRIO)\n"
                rag_context += "Use estas expressões e conectivos obrigatoriamente ao redigir:\n"
                rag_context += f"\n{glossary}\n"
        
        # ── ETAPA 2: Golden Sample / Caso Espelho (gabarito estrutural específico) ──
        retriever, _ = process_templates(template_files, api_key)
        
        if not retriever: return rag_context
        
        relevant_docs = retriever.invoke(text[:6000])
        
        if relevant_docs:
            mirror_doc = relevant_docs[0]
            other_docs = relevant_docs[1:]
            
            rag_context += "\n\n## 💎 CASO ESPELHO (GOLDEN SAMPLE - GABARITO ESTRUTURAL)\n"
            rag_context += f"⚠️ INSTRUÇÃO DE CLONAGEM ESTRUTURAL: O caso abaixo ({mirror_doc.metadata.get('source')}) é o seu GABARITO.\n"
            rag_context += "1. Copie a estrutura de tópicos (titulação, numeração).\n"
            rag_context += "2. Use os mesmos jargões e frases de transição do Glossário acima.\n"
            rag_context += "3. Se for o mesmo assunto, adapte apenas os fatos e nomes, mantendo a fundamentação jurídica.\n"
            rag_context += f"\n--- INÍCIO DO CASO ESPELHO ---\n{mirror_doc.page_content}\n--- FIM DO CASO ESPELHO ---\n"
            
            if other_docs:
                rag_context += "\n## OUTROS MODELOS DE REFERÊNCIA (CONTEXTO ADICIONAL)\n"
                for i, doc in enumerate(other_docs):
                    rag_context += f"\n[MODELO SECUNDÁRIO {i+2} - {doc.metadata.get('source')}]:\n{doc.page_content[:3000]}...\n"
                    
        return rag_context
    except Exception as e:
        print(f"Erro no retrieve_mirror_context: {e}")
        return ""


def run_standard_orchestration(text: str, main_llm_config: dict, style_llm_config: dict, status_callback=None, template_files=None, google_key=None, outline=None, style_prompt=None):
    """
    Pipeline Padrão (V1) FLEXÍVEL.
    Suporta qualquer LLM para Analista Principal e Analista de Estilo.
    """
    # Config keys: {'provider': str, 'model': str, 'key': str}
    
    def update(msg):
        if status_callback: status_callback(msg)

    try:
        # Instancia LLMs via Azure AI Foundry
        main_llm = get_llm("gpt-5.3-chat", temperature=0.2)
        style_llm = get_llm("gpt-5.3-chat", temperature=0.3)
    except Exception as e:
        return {"final_report": f"Erro na inicialização dos modelos: {str(e)}", "steps": {}}

    # PROCESSAMENTO DE TEMPLATES (RAG + DOSSIÊ FORENSE)
    rag_context = ""
    style_report = None
    style_dossier = None
    
    if template_files:
        # 1. Gerar Dossiê de Estilo Forense (cacheado) — usa Azure OpenAI internamente
        update("🧬 Gerando Dossiê de Identidade Decisional (5 Pilares)...")
        style_dossier = generate_style_dossier(template_files, None)
        style_report = style_dossier.get('dossier') if style_dossier else None
        
        # 2. Retrieve Mirror Context (agora com dossiê integrado)
        update("📚 Localizando Caso Espelho (Golden Sample)...")
        rag_context = retrieve_mirror_context(text, None, template_files, style_dossier=style_dossier)

    update(f"🧠 Iniciando Análise Profunda ({main_llm_config['model']})...")

    # 1. ANÁLISE INTEGRAL (MÉRITO/MINUTA)
    update("⚖️ Fase 1: Análise Integral e Minutagem (Analista Sênior)...")
    
    # --- LOAD KNOWLEDGE BASE (V4.5 Logic) ---
    kb_text = ""
    try:
        # NOTE: Mantendo Knowledge Base, mas REMOVENDO a substituição forçada de Prompt V3.
        # Queremos usar PROMPT_CLAUDE_INTEGRAL (JSON Mode)
        base_path = "data/knowledge_base"
        files_map = {
            "ARQUIVO A (Sobrestamento) 30.10.2025.txt": "ARQUIVO A (SOBRESTAMENTOS)",
            "ARQUIVO B (Súmulas) - 30.12.2025.txt": "ARQUIVO B (SÚMULAS)",
            "ARQUIVO C (Qualificados) - 30.12.2025.txt": "ARQUIVO C (QUALIFICADOS)"
        }
        for fname, label in files_map.items():
            fpath = os.path.join(base_path, fname)
            if os.path.exists(fpath):
                try:
                    with open(fpath, "r", encoding="utf-8") as f:
                        content = f.read()
                except UnicodeDecodeError:
                    with open(fpath, "r", encoding="latin-1") as f:
                        content = f.read()
                if content.strip():
                    kb_text += f"\n=== {label} ===\n{content}\n"
        
        # GARANTE USO DO PROMPT V1 OTIMIZADO (JSON)
        final_prompt_integral = PROMPT_CLAUDE_INTEGRAL
        if kb_text:
            final_prompt_integral += f"\n\n## 6. BASE DE CONHECIMENTO VINCULANTE (CARREGADA)\n{kb_text}"
            
    except Exception as e:
        final_prompt_integral = PROMPT_CLAUDE_INTEGRAL 
        print(f"Erro KB ou Prompt: {e}")

    # Injeta contexto RAG (Estilo)
    if rag_context:
        final_prompt_integral += rag_context
        
    # INJEÇÃO DO OUTLINE (PLANEJAMENTO)
    if outline:
        final_prompt_integral += f"\n\n## 📋 ESQUELETO LÓGICO (PLANEJAMENTO OBRIGATÓRIO)\nSiga estritamente esta estrutura para redigir a decisão:\n{outline}"
        
    # INJEÇÃO DE ESTILO (FEW-SHOT)
    # Se style_prompt for um FewShotTemplate, precisamos formatar.
    # Por simplicidade, se style_prompt existir, extraímos o texto dos exemplos
    few_shot_text = ""
    if style_prompt:
         try:
             # Formatação manual rápida dos exemplos para injetar no system
             for msg in style_prompt.format_messages(page_content=""):
                 few_shot_text += f"\nExemplo: {msg.content}\n"
             if few_shot_text:
                 final_prompt_integral += f"\n\n## 🎭 CLONAGEM DE ESTILO (RAG DINÂMICO)\nEscreva NO MESMO TOM destes exemplos:\n{few_shot_text}"
         except Exception as e:
             print(f"Erro ao formatar Style Prompt: {e}")

    integral_messages = [
        SystemMessage(content=final_prompt_integral),
        HumanMessage(content=f"Realize a ANÁLISE INTEGRAL E MINUTAGEM deste processo:\n\n[AUTOS DO PROCESSO]: {text[:200000]}") 
    ]
    integral_response = safe_content(main_llm.invoke(integral_messages))
    
    # --- REFLEXION LOOP (ACTIVE AUDITOR) ---
    update("🛡️ Rodando Auditoria Ativa (Verificando Alucinações)...")
    # run_reflexion_loop agora usa get_llm() internamente (Azure)
    final_output, audit_log = run_reflexion_loop(integral_response, text, None)
    
    return {
        "final_report": final_output,
        "auditor_dashboard": audit_log, 
        "style_report": style_report,
        "steps": {
            "integral": final_output
        }
    }

def run_ensemble_orchestration(text: str, keys: dict, status_callback=None, template_files=None):
    """
    V2: LINEAR ENSEMBLE PIPELINE (A "Linha de Montagem").
    Pipeline determinístico onde cada modelo faz uma parte específica.
    
    Fluxo:
    1. Gemini 3 Pro -> Extração de Fatos e Análise Formal (Input Massivo).
    2. DeepSeek R1 -> Análise Material/Mérito e Lógica Jurídica (Reasoning).
    3. Claude 4.6 Sonnet -> Redação Final (Minuta) com base nos insumos.
    """
    def update(msg):
         if status_callback: status_callback(msg)

    # 1. Setup Models
    try:
        # Todos os modelos via Azure AI Foundry (Claude Sonnet 4.6)
        analista_fatos = get_llm("gpt-5.3-chat", temperature=0.1)

        juiz_logico = get_llm("gpt-5.3-chat", temperature=0.3)
              
        redator_final = get_llm("gpt-5.3-chat", temperature=0.2)
             
    except Exception as e:
        return {"final_report": f"Erro ao inicializar Banca Digital: {e}", "steps": {}}

    logs = {}
    
    # MIRROR STRATEGY FOR V2 (com Dossiê Forense)
    rag_context = ""
    style_dossier = None
    if template_files:
        update("🧬 (V2) Gerando Dossiê de Identidade Decisional...")
        style_dossier = generate_style_dossier(template_files, keys['google'])
        update("📚 (V2) Localizando Caso Espelho...")
        rag_context = retrieve_mirror_context(text, keys['google'], template_files, style_dossier=style_dossier)

    # === FASE 1: EXTRAÇÃO E TRIAGEM (GEMINI) ===
    update("🕵️‍♂️ Fase 1: Gemini analisando Fatos e Requisitos Formais...")
    
    # Prompt de Fatos
    msg_fatos = [SystemMessage(content=PROMPT_FATOS), HumanMessage(content=f"Autos:\n{text[:150000]}")]
    res_fatos = safe_content(analista_fatos.invoke(msg_fatos))
    logs['fatos'] = res_fatos
    
    # Prompt Formal
    msg_formal = [SystemMessage(content=PROMPT_ANALISE_FORMAL), HumanMessage(content=f"Autos:\n{text[:100000]}")] # Menos contexto ok
    res_formal = safe_content(analista_fatos.invoke(msg_formal))
    logs['analise_formal'] = res_formal
    
    # === FASE 2: RACIOCÍNIO JURÍDICO (DEEPSEEK) ===
    # === FASE 2: RACIOCÍNIO JURÍDICO (DEEPSEEK) ===
    update("🧠 Fase 2: DeepSeek deliberando sobre o Mérito (Reasoning)...")
    
    # Monta o contexto para o Juiz
    contexto_juiz = f"""
    [RESUMO DOS FATOS]:
    {res_fatos}
    
    [TRIAGEM FORMAL]:
    {res_formal}
    
    [TRECHOS RELEVANTES DOS AUTOS]:
    {text[:50000]} 
    """
    
    # Prepare Mirror Context if available
    final_style_guide = keys.get('style_guide', "")
    if rag_context:
        final_style_guide += rag_context

    msg_material = PROMPT_JUIZ_DEEPSEEK.format(
        fatos_texto=res_fatos,
        formal_json=res_formal,
        style_guide=final_style_guide or "Estilo Padrão (Sem guia específico)."
    )
    
    # Use Invoke
    res_material = safe_content(juiz_logico.invoke([HumanMessage(content=contexto_juiz), SystemMessage(content=msg_material)]))
    logs['analise_material'] = res_material
    
    # === FASE 3: REDAÇÃO DE MINUTA (CLAUDE) ===
    update("✍️ Fase 3: Claude redigindo a Minuta Final (Sentença)...")
    
    msg_redator = PROMPT_REDATOR_CLAUDE.format(
        verdict_outline=res_material,
        style_guide=final_style_guide or "Estilo Padrão (Sem guia específico)."
    )
    
    res_final = safe_content(redator_final.invoke([HumanMessage(content=msg_redator)]))
    logs['minuta_final'] = res_final
    
    # === FASE 4: AUDITORIA FINAL (GPT-4o) ===
    final_output = res_final
    audit_log = "Auditoria GPT ignorada (Sem chave ou desativado)"

    # Check for OpenAI key availability
    if keys.get('openai') and HAS_OPENAI:
        update("🛡️ Fase 4: GPT-4o Auditando (Anti-Alucinação)...")
        try:
            auditor_gpt = ChatOpenAI(model="gpt-4o", api_key=keys['openai'], temperature=0.0)
            msg_audit = [
                SystemMessage(content=PROMPT_AUDITOR_GPT),
                HumanMessage(content=f"MINUTA PARA REVISÃO:\n{res_final}\n\nAUTOS:\n{text[:20000]}")
            ]
            audit_resp = safe_content(auditor_gpt.invoke(msg_audit))
            logs['auditoria_gpt'] = audit_resp
            
            if "ERRO:" in audit_resp or "REPROVADO" in audit_resp:
                audit_log = f"⚠️ ALERTA DO AUDITOR:\n{audit_resp}"
            else:
                audit_log = "✅ Aprovado pelo GPT-4o."
                
        except Exception as e:
            audit_log = f"Erro na auditoria GPT: {e}"

    return {
        "final_report": final_output,
        "auditor_dashboard": audit_log,
        "style_report": "Ensemble Assembly Line (Sem Style Guide específico)",
        "steps": logs
    }

# ── Template Store: Supabase-backed + in-memory cache ─────────────────────────
# Per-user: {user_id: [{"text": str, "metadata": dict}, ...]}
# Primary: Supabase REST API (persistent across deploys)
# Fallback: Local JSON files (for local development without Supabase)
# Cache: _template_store dict in memory (for fast TF-IDF queries)

import requests as _requests_lib

_template_store: dict[str, list[dict]] = {}
_TEMPLATE_STORE_BASE = os.path.abspath(os.getenv("CHROMA_DB_PATH", "./chroma_db_rag"))

# Ensure base directory exists (for local fallback)
os.makedirs(_TEMPLATE_STORE_BASE, exist_ok=True)

# Supabase config (reuse from environment)
_SUPA_URL = os.getenv("SUPABASE_URL", "") or os.getenv("VITE_SUPABASE_URL", "")
_SUPA_ANON_KEY = os.getenv("SUPABASE_ANON_KEY", "") or os.getenv("VITE_SUPABASE_ANON_KEY", "")
_SUPA_TABLE = "user_templates"

print(f"📂 Template store: {'Supabase (' + _SUPA_URL[:30] + '...)' if _SUPA_URL else 'Local JSON (' + _TEMPLATE_STORE_BASE + ')'}")


def _supa_headers(token: str = "") -> dict:
    """Build headers for Supabase REST API calls."""
    h = {
        "apikey": _SUPA_ANON_KEY,
        "Content-Type": "application/json",
        "Prefer": "return=representation",
    }
    if token:
        h["Authorization"] = f"Bearer {token}"
    return h


def _get_template_store_path(user_id: str = "default") -> str:
    """Returns user-specific template store path (local fallback)."""
    return os.path.join(_TEMPLATE_STORE_BASE, user_id, "templates.json")

# Legacy path for backward compatibility
_TEMPLATE_STORE_PATH = os.path.join(_TEMPLATE_STORE_BASE, "templates.json")


class SimpleRetriever:
    """Retriever using TF-IDF cosine similarity. Falls back to keyword overlap if sklearn unavailable."""
    
    # Portuguese stopwords for legal text
    _STOP_WORDS_PT = {
        'a', 'o', 'e', 'de', 'do', 'da', 'dos', 'das', 'em', 'no', 'na', 'nos', 'nas',
        'um', 'uma', 'uns', 'umas', 'ao', 'aos', 'à', 'às', 'por', 'para', 'com', 'sem',
        'que', 'se', 'não', 'mais', 'como', 'mas', 'ou', 'quando', 'muito', 'já', 'também',
        'só', 'seu', 'sua', 'seus', 'suas', 'esse', 'essa', 'esses', 'essas', 'este', 'esta',
        'estes', 'estas', 'isso', 'isto', 'aquele', 'aquela', 'aqueles', 'aquelas', 'aquilo',
        'ele', 'ela', 'eles', 'elas', 'meu', 'minha', 'nós', 'vós', 'ter', 'ser', 'estar',
        'foi', 'são', 'será', 'seria', 'tem', 'tinha', 'entre', 'sobre', 'após', 'até',
        'pelo', 'pela', 'pelos', 'pelas', 'qual', 'quais', 'onde', 'quem', 'porque',
        'ainda', 'mesmo', 'pode', 'deve', 'assim', 'bem', 'todo', 'toda', 'todos', 'todas',
        'cada', 'outro', 'outra', 'outros', 'outras', 'parte', 'forma', 'conforme',
    }
    
    def __init__(self, docs: list[dict], k: int = 5):
        self.docs = docs
        self.k = k
        self._tfidf_matrix = None
        self._vectorizer = None
        self._build_index()
    
    def _build_index(self):
        """Pre-compute TF-IDF matrix if sklearn is available."""
        if not self.docs:
            return
        try:
            from sklearn.feature_extraction.text import TfidfVectorizer
            self._vectorizer = TfidfVectorizer(
                max_features=8000,
                stop_words=list(self._STOP_WORDS_PT),
                ngram_range=(1, 2),
                sublinear_tf=True,
            )
            corpus = [d["text"] for d in self.docs]
            self._tfidf_matrix = self._vectorizer.fit_transform(corpus)
            print(f"📊 TF-IDF index built: {len(corpus)} docs, {self._tfidf_matrix.shape[1]} features")
        except ImportError:
            print("⚠️ sklearn não disponível. Usando fallback keyword overlap.")
            self._vectorizer = None
            self._tfidf_matrix = None
    
    def invoke(self, query: str) -> list:
        """Return top-k documents ranked by TF-IDF cosine similarity (or keyword fallback)."""
        if not self.docs:
            return []
        from langchain_core.documents import Document
        scored_docs = self._score(query)
        results = []
        for score, doc_dict in scored_docs[:self.k]:
            doc = Document(
                page_content=doc_dict["text"],
                metadata={**doc_dict.get("metadata", {}), "relevance_score": round(score, 4)}
            )
            results.append(doc)
        return results
    
    def invoke_scored(self, query: str) -> list[tuple]:
        """Return top-k as (score, Document) tuples for quality-aware injection."""
        if not self.docs:
            return []
        from langchain_core.documents import Document
        scored_docs = self._score(query)
        results = []
        for score, doc_dict in scored_docs[:self.k]:
            doc = Document(
                page_content=doc_dict["text"],
                metadata={**doc_dict.get("metadata", {}), "relevance_score": round(score, 4)}
            )
            results.append((round(score, 4), doc))
        return results
    
    def _score(self, query: str) -> list[tuple]:
        if self._tfidf_matrix is not None and self._vectorizer is not None:
            return self._score_tfidf(query)
        return self._score_keyword(query)
    
    def _score_tfidf(self, query: str) -> list[tuple]:
        from sklearn.metrics.pairwise import cosine_similarity
        query_vec = self._vectorizer.transform([query])
        scores = cosine_similarity(query_vec, self._tfidf_matrix).flatten()
        scored = [(scores[i], self.docs[i]) for i in range(len(self.docs))]
        scored.sort(key=lambda x: x[0], reverse=True)
        return scored
    
    def _score_keyword(self, query: str) -> list[tuple]:
        query_words = set(query.lower().split()) - self._STOP_WORDS_PT
        scored = []
        for doc_dict in self.docs:
            doc_words = set(doc_dict["text"].lower().split()) - self._STOP_WORDS_PT
            if not query_words or not doc_words:
                scored.append((0.0, doc_dict))
                continue
            overlap = len(query_words & doc_words)
            score = overlap / max(len(query_words), 1)
            scored.append((score, doc_dict))
        scored.sort(key=lambda x: x[0], reverse=True)
        return scored


# ── Theme extraction for intelligent RAG queries ──────────────────────────

def extract_legal_themes(process_text: str) -> str:
    """
    Extracts 3-5 core legal themes from a process text using a fast LLM.
    Used as query for model search instead of raw text.
    Returns a concise theme string for retriever queries.
    """
    try:
        llm = get_llm("gpt-5.4-mini", temperature=0.0, max_tokens=200)
        messages = [
            SystemMessage(content=(
                "Você é um classificador jurídico. Dado o texto de um processo, extraia os 3 a 5 temas jurídicos centrais "
                "em palavras-chave objetivas, separadas por vírgula. Responda APENAS com as palavras-chave, sem explicação.\n"
                "Exemplos de saída: 'dano moral, relação de consumo, atraso de voo, indenização'\n"
                "'contrato de compra e venda, inadimplemento, rescisão contratual, restituição de valores'"
            )),
            HumanMessage(content=process_text[:4000])
        ]
        response = llm.invoke(messages)
        themes = safe_content(response).strip()
        print(f"🏷️ Temas jurídicos extraídos: {themes}")
        return themes
    except Exception as e:
        print(f"⚠️ Erro na extração de temas: {e}")
        return process_text[:500]


# ── Per-user style dossier persistence ──────────────────────────────────

def _get_dossier_path(user_id: str) -> str:
    """Get path for user's style dossier JSON file."""
    return os.path.join(_TEMPLATE_STORE_BASE, f"user_{user_id}", "style_dossier.json")


def save_style_dossier(user_id: str, dossier: dict):
    """Persist style dossier to disk for a user."""
    try:
        path = _get_dossier_path(user_id)
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(dossier, f, ensure_ascii=False)
        _style_dossier_cache[f"user:{user_id}"] = dossier
        print(f"💾 Dossiê de estilo salvo para user {user_id[:8]} ({len(dossier.get('dossier', ''))} chars)")
    except Exception as e:
        print(f"⚠️ Erro ao salvar dossiê: {e}")


def load_style_dossier(user_id: str) -> dict:
    """Load style dossier from memory cache or disk."""
    cache_key = f"user:{user_id}"
    if cache_key in _style_dossier_cache:
        return _style_dossier_cache[cache_key]
    try:
        path = _get_dossier_path(user_id)
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                dossier = json.load(f)
            _style_dossier_cache[cache_key] = dossier
            return dossier
    except Exception as e:
        print(f"⚠️ Erro ao carregar dossiê do disco: {e}")
    return None


# ── Sentence template for fallback (injected only when no RAG models) ──

TEMPLATE_SENTENCA_PADRAO = """
## TEMPLATE PADRÃO DE SENTENÇA (Art. 489 CPC)
Use esta estrutura APENAS se não houver modelos do magistrado disponíveis.

**RELATÓRIO**
Trata-se de Ação [Natureza] ajuizada por [Autor] em face de [Réu].
Narra a parte autora, em síntese, que [causa de pedir]. Requer [pedidos]. Juntou documentos (ID X).
[Se liminar]: A tutela provisória foi [deferida/indeferida] em ID X.
Citado(a) (ID X), o(a) réu apresentou contestação (ID Y), arguindo [preliminares]. No mérito, sustenta que [defesa].
[Se réplica]: Houve réplica em ID Z.
É o relatório. Decido.

**FUNDAMENTAÇÃO**

**I. Questões Processuais e Preliminares**
[Análise de cada preliminar com acolhimento/rejeição fundamentada]

**II. Prejudiciais de Mérito**
[Se prescrição/decadência: análise cronológica]

**III. Mérito**
[Desenvolvimento analítico de cada ponto controvertido com fundamentação legal]

**DISPOSITIVO**
Ante o exposto, e por tudo mais que dos autos consta:
**JULGO [PROCEDENTE/IMPROCEDENTE/PARCIALMENTE PROCEDENTE]** o(s) pedido(s), com resolução de mérito (art. 487, I, CPC), para:
1. [Condenação/Obrigação/Rejeição]

**Sucumbência:**
Condeno a parte [vencida] ao pagamento das custas e honorários advocatícios, fixados em [%] sobre o valor [da condenação/da causa] (art. 85, §2º, CPC).
[Se JG]: Suspendo a exigibilidade (art. 98, §3º, CPC).

P.R.I.
"""


# ── Supabase-backed save/load (primary) with local JSON fallback ─────────
import re as _re

_UUID_RE = _re.compile(r'^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$', _re.I)

def _is_valid_uuid(value: str) -> bool:
    """Return True only if value is a well-formed UUID (required by Supabase UUID column)."""
    return bool(value and _UUID_RE.match(value))


def _save_template_store(user_id: str = "default", token: str = ""):
    """Persist template store: try Supabase first, then local JSON fallback."""
    user_data = _template_store.get(user_id, [])
    
    # Strategy 1: Supabase REST (only when user_id is a well-formed UUID)
    if _SUPA_URL and _SUPA_ANON_KEY and token and _is_valid_uuid(user_id):
        try:
            # Delete existing chunks for this user, then insert new ones
            del_url = f"{_SUPA_URL.rstrip('/')}/rest/v1/{_SUPA_TABLE}?user_id=eq.{user_id}"
            _requests_lib.delete(del_url, headers=_supa_headers(token), timeout=10)
            
            if user_data:
                # Batch insert (Supabase accepts array of objects)
                rows = []
                for chunk in user_data:
                    rows.append({
                        "user_id": user_id,
                        "source": chunk.get("metadata", {}).get("source", "unknown"),
                        "text": chunk["text"],
                        "upload_date": chunk.get("metadata", {}).get("upload_date", None),
                    })
                
                # Insert in batches of 100 to avoid payload limits
                BATCH_SIZE = 100
                for i in range(0, len(rows), BATCH_SIZE):
                    batch = rows[i:i + BATCH_SIZE]
                    ins_url = f"{_SUPA_URL.rstrip('/')}/rest/v1/{_SUPA_TABLE}"
                    resp = _requests_lib.post(ins_url, json=batch, headers=_supa_headers(token), timeout=15)
                    if not resp.ok:
                        print(f"⚠️ Supabase insert batch {i//BATCH_SIZE+1} failed ({resp.status_code}): {resp.text[:200]}")
                        raise Exception(f"Supabase insert failed: {resp.status_code}")
                
                print(f"💾 Templates salvos no Supabase: {len(user_data)} chunks (user {user_id[:8]})")
                return
            else:
                print(f"💾 Templates limpos no Supabase (user {user_id[:8]})")
                return
        except Exception as e:
            print(f"⚠️ Supabase save failed, falling back to local: {e}")
    
    # Strategy 2: Local JSON fallback
    try:
        path = _get_template_store_path(user_id)
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(user_data, f, ensure_ascii=False)
        print(f"💾 Templates salvos localmente: {path} ({len(user_data)} chunks)")
    except Exception as e:
        print(f"❌ ERRO ao salvar templates: {e}")


def _load_template_store(user_id: str = "default", token: str = ""):
    """Load template store: try Supabase first, then local JSON fallback."""
    global _template_store
    
    # Strategy 1: Supabase REST (only when user_id is a well-formed UUID)
    if _SUPA_URL and _SUPA_ANON_KEY and token and _is_valid_uuid(user_id):
        try:
            url = f"{_SUPA_URL.rstrip('/')}/rest/v1/{_SUPA_TABLE}?user_id=eq.{user_id}&select=source,text,upload_date"
            resp = _requests_lib.get(url, headers=_supa_headers(token), timeout=10)
            if resp.ok:
                rows = resp.json()
                if rows:
                    _template_store[user_id] = [
                        {"text": r["text"], "metadata": {"source": r["source"], "upload_date": r.get("upload_date", "")}}
                        for r in rows
                    ]
                    print(f"✅ Templates carregados do Supabase: {len(rows)} chunks (user {user_id[:8]})")
                    return
                else:
                    print(f"📭 Nenhum template no Supabase para user {user_id[:8]}, tentando JSON local...")
            else:
                print(f"⚠️ Supabase load failed ({resp.status_code}): {resp.text[:200]}")
        except Exception as e:
            print(f"⚠️ Supabase load failed, trying local: {e}")
    
    # Strategy 2: Local JSON fallback
    try:
        path = _get_template_store_path(user_id)
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                _template_store[user_id] = json.load(f)
            print(f"✅ Templates carregados do disco: {len(_template_store[user_id])} chunks (user {user_id[:8]})")
            return
        if user_id != "default" and os.path.exists(_TEMPLATE_STORE_PATH):
            with open(_TEMPLATE_STORE_PATH, "r", encoding="utf-8") as f:
                legacy_data = json.load(f)
            if legacy_data:
                print(f"🔄 Migrando templates legados para user {user_id[:8]}...")
                _template_store[user_id] = legacy_data
                _save_template_store(user_id, token)
                return
        _template_store[user_id] = []
    except Exception as e:
        print(f"❌ ERRO ao carregar templates: {e}")
        _template_store[user_id] = []


def _delete_templates_supabase(user_id: str, source: str = "", token: str = "") -> bool:
    """Delete templates from Supabase. If source is given, deletes only that source. Returns True on success."""
    if not (_SUPA_URL and _SUPA_ANON_KEY and token and _is_valid_uuid(user_id)):
        return False
    try:
        url = f"{_SUPA_URL.rstrip('/')}/rest/v1/{_SUPA_TABLE}?user_id=eq.{user_id}"
        if source:
            url += f"&source=eq.{source}"
        resp = _requests_lib.delete(url, headers=_supa_headers(token), timeout=10)
        if resp.ok:
            print(f"🗑️ Templates deletados do Supabase (user {user_id[:8]}, source={source or 'ALL'})")
            return True
        print(f"⚠️ Supabase delete failed ({resp.status_code}): {resp.text[:200]}")
    except Exception as e:
        print(f"⚠️ Supabase delete error: {e}")
    return False


def process_templates(files, api_key, collection_name="rag_templates_persistent", user_id="default", token=""):
    """
    Processa arquivos de template (PDF/DOCX/TXT) e cria um retriever.
    100% local chunking. Persists to Supabase (or local JSON fallback).
    """
    import time as _time
    t0 = _time.time()
    global _template_store
    
    documents = []
    splitter = RecursiveCharacterTextSplitter(chunk_size=4000, chunk_overlap=200)

    for file in files:
        if hasattr(file, 'seek'):
            file.seek(0)
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file.name.split('.')[-1]}") as tmp:
            tmp.write(file.read())
            tmp_path = tmp.name
        
        try:
            text = ""
            if file.name.endswith(".pdf"):
                reader = pypdf.PdfReader(tmp_path)
                for page in reader.pages:
                    text += (page.extract_text() or "") + "\n"
            elif file.name.endswith(".docx"):
                doc = docx.Document(tmp_path)
                text = "\n".join([p.text for p in doc.paragraphs])
            else:
                try:
                    with open(tmp_path, "r", encoding="utf-8") as f: text = f.read()
                except UnicodeDecodeError:
                     with open(tmp_path, "r", encoding="latin-1") as f: text = f.read()
            
            import datetime as _dt
            doc_chunks = splitter.create_documents([text], metadatas=[{"source": file.name, "upload_date": _dt.datetime.now().isoformat()}])
            documents.extend(doc_chunks)
        finally:
            os.remove(tmp_path)
    
    if not documents:
        return None, []

    # Store in memory cache
    _template_store[user_id] = [
        {"text": doc.page_content, "metadata": doc.metadata}
        for doc in documents
    ]
    # Persist (Supabase or local JSON)
    _save_template_store(user_id, token)
    
    print(f"✅ RAG indexado (user {user_id[:8]}): {len(documents)} chunks em {_time.time()-t0:.1f}s")
    return SimpleRetriever(_template_store[user_id]), documents


def load_persistent_rag(api_key=None, collection_name="rag_templates_persistent", user_id="default", token=""):
    """Load persisted templates for RAG retrieval."""
    global _template_store
    try:
        if user_id not in _template_store or not _template_store[user_id]:
            _load_template_store(user_id, token)
        user_store = _template_store.get(user_id, [])
        if user_store:
            print(f"RAG Persistente carregado (user {user_id[:8]}): {len(user_store)} chunks.")
            return SimpleRetriever(user_store)
    except Exception as e:
        print(f"Erro ao carregar RAG persistente para user {user_id[:8]}: {e}")
    return None

def list_templates(user_id: str = "default", token: str = "") -> list[dict]:
    """List all template files for a user, grouped by source filename."""
    global _template_store
    if user_id not in _template_store or not _template_store[user_id]:
        _load_template_store(user_id, token)
    user_store = _template_store.get(user_id, [])
    
    sources: dict[str, dict] = {}
    for chunk in user_store:
        source = chunk.get("metadata", {}).get("source", "desconhecido")
        if source not in sources:
            sources[source] = {
                "filename": source,
                "chunk_count": 0,
                "upload_date": chunk.get("metadata", {}).get("upload_date", None),
                "total_chars": 0,
            }
        sources[source]["chunk_count"] += 1
        sources[source]["total_chars"] += len(chunk.get("text", ""))
    
    return list(sources.values())


def delete_template_by_source(user_id: str, source_name: str, token: str = "") -> int:
    """Delete all chunks from a specific source file. Returns count of removed chunks."""
    global _template_store
    if user_id not in _template_store or not _template_store[user_id]:
        _load_template_store(user_id, token)
    
    user_store = _template_store.get(user_id, [])
    original_count = len(user_store)
    
    _template_store[user_id] = [
        chunk for chunk in user_store
        if chunk.get("metadata", {}).get("source", "") != source_name
    ]
    
    removed = original_count - len(_template_store[user_id])
    if removed > 0:
        # Supabase: delete specific source
        _delete_templates_supabase(user_id, source_name, token)
        # Local fallback
        _save_template_store(user_id, token)
    
    return removed


def search_templates(user_id: str, query: str, k: int = 5, token: str = "") -> list[dict]:
    """Search user's templates using TF-IDF retriever. Returns top-k results."""
    global _template_store
    if user_id not in _template_store or not _template_store[user_id]:
        _load_template_store(user_id, token)
    user_store = _template_store.get(user_id, [])
    if not user_store:
        return []
    
    retriever = SimpleRetriever(user_store, k=k)
    results = retriever.invoke(query)
    
    return [
        {
            "text": doc.page_content[:1000],
            "source": doc.metadata.get("source", "?"),
            "full_length": len(doc.page_content),
        }
        for doc in results
    ]


def generate_style_report(documents, api_key):
    """
    Usa um modelo rápido para ler os templates e criar um perfil de estilo.
    Migrado para Azure OpenAI (gpt-5.4-mini).
    """
    try:
        llm_flash = get_llm("gpt-5.4-mini", temperature=0.3)
        
        
        # Concatena amostras dos documentos (Random Sampling)
        sample_text = ""
        # Seleciona de 3 a 5 chunks aleatórios para ter variabilidade
        num_samples = min(5, len(documents))
        if num_samples > 0:
            selected_docs = random.sample(documents, num_samples)
            for doc in selected_docs:
                sample_text += f"\n--- AMOSTRA ({doc.metadata.get('source')}): ---\n{doc.page_content[:5000]}\n"
            
        messages = [
            SystemMessage(content=PROMPT_STYLE_ANALYZER),
            HumanMessage(content=f"Aqui estão amostras de decisões do magistrado. Crie o Dossiê de Estilo:\n{sample_text}")
        ]
        
        response = llm_flash.invoke(messages)
        content = safe_content(response)
            
        return clean_text(content)
    except Exception as e:
        return f"Erro ao gerar perfil de estilo: {str(e)}"

# OLD run_gemini_orchestration removed/replaced by run_standard_orchestration

def process_batch(files, api_key):
    """
    Processa múltiplos arquivos (PDF/DOCX) para o X-Ray.
    Retorna uma lista de strings (textos extraídos).
    """
    processed_texts = []
    
    for file in files:
        # Reutiliza a lógica de extração salvando em temp
        suffix = os.path.splitext(file.name)[1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_file.write(file.read())
            tmp_path = tmp_file.name
        
        try:
           # Extração simplificada, sem vetorização individual
           text_content = ""
           if suffix == ".pdf":
               loader = PyPDFLoader(tmp_path)
               docs = loader.load()
               text_content = "\n".join([d.page_content for d in docs])
           elif suffix == ".docx":
               from langchain_community.document_loaders import Docx2txtLoader
               loader = Docx2txtLoader(tmp_path)
               docs = loader.load()
               text_content = "\n".join([d.page_content for d in docs])
           elif suffix == ".txt":
               from langchain_community.document_loaders import TextLoader
               loader = TextLoader(tmp_path)
               docs = loader.load()
               text_content = "\n".join([d.page_content for d in docs])
           
           if text_content:
               processed_texts.append(f"--- PROCESSO: {file.name} ---\n{clean_text(text_content[:20000])}") # Limita chars por doc para caber no contexto
               
        except Exception as e:
            print(f"Erro ao ler {file.name}: {e}")
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
                
    return processed_texts

from prompts_claude import PROMPT_XRAY_MAP, PROMPT_XRAY_BATCH

def map_process_individual(text_content, filename, api_key=None):
    """
    ETAPA MAP: Analisa um único processo e retorna JSON estruturado.
    Usa GPT-4.1-mini (Azure) para rapidez e custo baixíssimo nesta triagem massiva paralela.
    Retries on 429 rate limit.
    """
    import time as _time

    MAX_RETRIES = 3
    BASE_DELAY = 10  # seconds

    for attempt in range(MAX_RETRIES + 1):
        try:
            llm = get_llm("gpt-5.4-mini", temperature=0.1)
            
            # Força strict JSON no prompt
            map_prompt = PROMPT_XRAY_MAP + "\n\nCRÍTICO: Retorne APENAS UM JSON (Strict JSON). Nenhuma palavra fora das chaves {}."
            
            messages = [
                SystemMessage(content=map_prompt),
                HumanMessage(content=f"Arquivo: {filename}\n\n{text_content[:20000]}")
            ]
            response = safe_content(llm.invoke(messages))
            
            # Limpa JSON
            cleaned = response.replace("```json", "").replace("```", "").strip()
            data = json.loads(cleaned)
            data["filename"] = filename # Garante que o nome do arquivo persista
            return data
            
        except Exception as e:
            error_str = str(e)
            if "429" in error_str and attempt < MAX_RETRIES:
                delay = BASE_DELAY * (2 ** attempt)  # 10s, 20s, 40s
                print(f"⚠️ Rate limit (429) no Map de {filename}, retry {attempt + 1}/{MAX_RETRIES} após {delay}s...")
                _time.sleep(delay)
                continue
            print(f"Falha total no Map de {filename} com gpt-5.4-mini (Azure). Erro: {e}")
            return {
                "filename": filename, 
                "error": f"Falha na leitura (GPT-4.1-mini). Err: {str(e)}", 
                "sintese_fatos": "Erro de leitura estruturada", 
                "tags_juridicas": ["ERRO"]
            }

def generate_batch_xray(files, api_key, template_files=None, progress_callback=None):
    """
    Gera o Raio-X da carteira usando estratégia MAP-REDUCE.
    1. MAP: Extrai metadados de cada processo individualmente (Paralelo via Gemini Flash).
    2. REDUCE: Envia lista de metadados para o Gemini Flash agrupar.
    progress_callback: optional callable(msg: str) to report progress.
    """
    def _progress(msg):
        if progress_callback:
            progress_callback(msg)
        print(f"📊 Raio-X: {msg}")

    try:
        # 1. PROCESSAMENTO DE TEXTO (Leitura)
        _progress("Extraindo texto dos arquivos...")
        raw_texts = []
        for file in files:
            suffix = os.path.splitext(file.name)[1].lower()
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
                tmp_file.write(file.read())
                tmp_path = tmp_file.name
            
            try:
                content = ""
                if suffix == ".pdf":
                    loader = PyPDFLoader(tmp_path)
                    docs = loader.load()
                    content = "\n".join([d.page_content for d in docs])
                elif suffix == ".docx":
                    from langchain_community.document_loaders import Docx2txtLoader
                    loader = Docx2txtLoader(tmp_path)
                    docs = loader.load()
                    content = "\n".join([d.page_content for d in docs])
                elif suffix == ".txt":
                    from langchain_community.document_loaders import TextLoader
                    try:
                        loader = TextLoader(tmp_path, encoding='utf-8')
                        content = loader.load()[0].page_content
                    except Exception:
                        loader = TextLoader(tmp_path, encoding='latin-1')
                        content = loader.load()[0].page_content
                
                if content:
                    raw_texts.append((file.name, clean_text(content)))
            except Exception as e:
                print(f"Erro lendo {file.name}: {e}")
            finally:
                if os.path.exists(tmp_path): os.remove(tmp_path)

        if not raw_texts:
            return {"error": "Nenhum texto extraído."}

        total_files = len(raw_texts)
        _progress(f"Texto extraído de {total_files} arquivos. Iniciando análise MAP...")

        # 2. ETAPA MAP (Execução com paralelismo otimizado para Azure S0)
        import time as _time
        mapped_data = []
        completed_count = 0

        with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
            futures = {}
            for i, (fname, text) in enumerate(raw_texts):
                if i > 0:
                    _time.sleep(2)  # Stagger para respeitar rate limit Azure S0
                future = executor.submit(map_process_individual, text, fname, api_key)
                futures[future] = fname
            for future in concurrent.futures.as_completed(futures):
                try:
                    res = future.result()
                    mapped_data.append(res)
                    completed_count += 1
                    _progress(f"Analisando arquivo {completed_count} de {total_files} ({futures[future]})...")
                except Exception as e:
                    completed_count += 1
                    print(f"Erro no map thread: {e}")

        _progress(f"MAP concluído ({len(mapped_data)} fichas). Consolidando clusters (REDUCE)...")

        # 3. ETAPA REDUCE (Clusterização)
        mapped_json_str = json.dumps(mapped_data, ensure_ascii=False, indent=2)
        
        # Cria dicionário de cache para retorno {filename: text}
        text_cache = {fname: text for fname, text in raw_texts}
        
        # Prepara Contexto de Modelos (Templates)
        models_context = ""
        if template_files:
            model_texts = process_batch(template_files, api_key)
            if model_texts:
                 models_context = "\n\n## MODELOS DE REFERÊNCIA DISPONÍVEIS:\n" + "\n".join(model_texts)
        
        human_msg = f"""
        Aqui estão as FICHAS TÉCNICAS dos processos processados individualmente.
        Agrupe-os e gere o relatório de Raio-X.
        
        [DADOS DOS PROCESSOS (JSON)]:
        {mapped_json_str}
        
        {models_context}
        """
        
        # Usa GPT-4.1-mini para agregação rápida (com retry)
        MAX_REDUCE_RETRIES = 3
        REDUCE_BASE_DELAY = 10
        content = None

        for reduce_attempt in range(MAX_REDUCE_RETRIES + 1):
            try:
                llm_reduce = get_llm("gpt-5.4-mini", temperature=0.1)
                
                # Força JSON
                reduce_prompt = PROMPT_XRAY_BATCH + "\n\nCRÍTICO: Retorne APENAS UM JSON VÁLIDO. Sem Markdown, sem formatação extra, inicie com { e termine com }."
                
                messages = [
                    SystemMessage(content=reduce_prompt),
                    HumanMessage(content=human_msg)
                ]
                
                response = safe_content(llm_reduce.invoke(messages))
                content = response
                break  # Success, exit retry loop
                
            except Exception as e:
                error_str = str(e)
                if "429" in error_str and reduce_attempt < MAX_REDUCE_RETRIES:
                    delay = REDUCE_BASE_DELAY * (2 ** reduce_attempt)
                    print(f"⚠️ Rate limit (429) no Reduce, retry {reduce_attempt + 1}/{MAX_REDUCE_RETRIES} após {delay}s...")
                    _time.sleep(delay)
                    continue
                print(f"Erro Crítico no Reduce (GPT-4.1-mini): {e}")
                return {"error": f"Erro na consolidação de dados. Detalhe: {str(e)}", "raw_content": ""}, text_cache

        _progress("Finalizando relatório...")

        
        # Garante que content é string (algumas versões retornam lista)
        if isinstance(content, list):
            # Se for lista de strings ou objetos com text, tenta converter
            try:
                content = "".join([str(c) for c in content])
            except Exception:
                content = str(content)
        elif not isinstance(content, str):
            content = str(content)
        
        # Limpeza do JSON
        try:
            # Tenta extrair bloco JSON delimitado por markdown
            json_match = re.search(r"```json\s*(.*?)```", content, re.DOTALL)
            if json_match:
                cleaned_json = json_match.group(1).strip()
            else:
                # Fallback: Tenta encontrar o maior bloco JSON possível ({...} ou [...])
                match = re.search(r"(\[.*\]|\{.*\})", content, re.DOTALL)
                if match:
                    cleaned_json = match.group(1).strip()
                else:
                    cleaned_json = content.replace("```json", "").replace("```", "").strip()

            try:
                return json.loads(cleaned_json), text_cache
            except json.JSONDecodeError:
                # Tenta corrigir JSON malformado (ex: aspas simples, trailing commas)
                 try:
                     import ast
                     # ast.literal_eval consegue parsear dicts python stringficados ({'key': 'val'})
                     repaired = ast.literal_eval(cleaned_json)
                     return repaired, text_cache
                 except Exception:
                     # Última tentativa: regex replace de aspas simples
                     try:
                         repaired_str = cleaned_json.replace("'", '"').replace("Mm.", "Mm").replace("Exa.", "Exa") # Hacks comuns
                         return json.loads(repaired_str), text_cache
                     except Exception:
                        pass
                 
                 # Se tudo falhar, retorna erro
                 return {"error": "Falha ao decodificar JSON do Reduce", "raw_content": content}, text_cache

        except Exception as inner_e:
             return {"error": f"Erro interno JSON: {str(inner_e)}", "raw_content": content}, text_cache
        
    except Exception as e:
        return {"error": f"Erro Geral no Pipeline: {str(e)}\n{traceback.format_exc()}"}, {}




def process_single_case_pipeline(pdf_bytes, filename, api_key, template_files=None, cached_text=None, mode="v1", keys=None, ocr_engine_choice="marker"):
    """
    Função Worker para processar um único caso completo.
    Suporta V1 (Gemini Only) e V2 (Hybrid Agents).
    """
    try:
        # 1. Extract Text
        if cached_text:
            text_content = cached_text
            clean_content = text_content # Já deve vir limpo do cache
        else:
            # Fallback para leitura de bytes se não tiver cache
            # Precisamos salvar bytes em temp file para loaders funcionarem
            suffix = os.path.splitext(filename)[1].lower()
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(pdf_bytes)
                tmp_path = tmp.name
            
            try:
                if suffix == ".pdf":
                    # Tentativa 1: Leitura de Texto Nativo
                    loader = PyPDFLoader(tmp_path)
                    docs = loader.load()
                    text_content = "\n".join([d.page_content for d in docs])
                    
                    # Tentativa 2: OCR Avançado (se texto vazio e OCR habilitado)
                    # Se tiver menos de 100 caracteres
                    if len(text_content.strip()) < 100 and HAS_OCR:
                        print(f"⚠️ Texto insuficiente ({len(text_content)} chars) em {filename}. Iniciando OCR Marker...")
                        try:
                            # Chama o Marker engine
                            from ocr_engine import get_marker_engine
                            marker_eng = get_marker_engine()
                            ocr_text = marker_eng.process_pdf(tmp_path) if marker_eng else ""
                            
                            # Se OCR retornou algo razoável, usa
                            if len(ocr_text) > len(text_content):
                                text_content = ocr_text
                                print(f"✅ OCR Avançado extraiu {len(text_content)} caracteres.")
                            elif "[ERRO]" in ocr_text:
                                print(f"Falha no OCR: {ocr_text}")
                                
                        except Exception as e:
                             print(f"Erro no OCR Pipeline: {e}")

                             
                elif suffix == ".docx":
                    from langchain_community.document_loaders import Docx2txtLoader
                    loader = Docx2txtLoader(tmp_path)
                    docs = loader.load()
                    text_content = "\n".join([d.page_content for d in docs])
                elif suffix == ".txt":
                    try:
                        with open(tmp_path, "r", encoding="utf-8") as f: text_content = f.read()
                    except UnicodeDecodeError:
                        with open(tmp_path, "r", encoding="latin-1") as f: text_content = f.read()
                else:
                    text_content = ""
            finally:
                if os.path.exists(tmp_path): os.remove(tmp_path)
                
            try:
                clean_content = clean_text(text_content)
                # Ensure it's a valid string, replacing bad chars if needed
                if isinstance(clean_content, bytes):
                    clean_content = clean_content.decode('utf-8', errors='replace')
            except Exception as e_clean:
                clean_content = "Erro de decodificação de texto."
                print(f"Erro no clean_text: {e_clean}")
        
        # --- RAPTOR INTEGRATION (LONG CONTEXT) ---
        # Se o texto for muito grande (> 150k chars), aciona o RAPTOR para resumir
        if len(clean_content) > 150000:
             print(f"🦖 RAPTOR ATIVADO: Texto Grande ({len(clean_content)} chars). Iniciando indexação hierárquica...")
             try:
                 # Decide key/provider
                 raptor_key = keys.get('openai') if keys and keys.get('openai') else (keys.get('google') if keys else api_key)
                 raptor_provider = "openai" if (keys and keys.get('openai')) else "google"
                 
                 raptor = RaptorEngine(api_key=raptor_key, provider=raptor_provider)
                 
                 # Gera Árvore de Resumos
                 tree_summary = raptor.build_tree(clean_content)
                 
                 # Substitui o texto original pela Árvore (que é bem menor e focada)
                 # Mantendo um prefixo identificando
                 clean_content = f" [MODO RAPTOR ATIVADO]\nO texto a seguir é um RESUMO HIERÁRQUICO do processo original.\n\n{tree_summary}"
                 print("✅ RAPTOR finalizado. Texto reduzido com sucesso.")
                 
             except Exception as e_raptor:
                 print(f"⚠️ Erro ao executar RAPTOR: {e_raptor}. Usando texto original truncado.")
        
        # --- AGENTIC WORKFLOW (PLANNER -> WRITER -> CRITIC) ---
        # Substitui a lógica manual anterior pelo Grafo do LangGraph
        
        final_draft = ""
        outline = None       # Inicialização para evitar NameError no fallback V1
        style_prompt = None  # Idem
        
        # Decide keys for agents
        agent_key = keys.get('openai') if keys and keys.get('openai') else (keys.get('google') if keys else api_key)
        agent_provider = "openai" if (keys and keys.get('openai')) else "google"

        # Verifica se o workflow está disponível
        if create_agent_workflow is not None:
            try:
                 print("🚀 Iniciando Workflow Agêntico (Planner -> Writer -> Critic)...")
                 app = create_agent_workflow()
                 
                 inputs = {
                     "facts": clean_content,
                     "api_key": agent_key,
                     "provider": agent_provider,
                     "revision_count": 0
                 }
                 
                 # Executa o Grafo
                 result_state = app.invoke(inputs)
                 final_draft = result_state.get("draft", "Erro na geração do draft.")
                 
                 print("✅ Workflow Completo com Sucesso!")
                 
            except Exception as e_workflow:
                 print(f"⚠️ Erro no Workflow Agêntico: {e_workflow}. Caindo para pipeline legado.")
                 # Fallback logic could be here if needed, but for now allow flow to continue to standard orchestration if draft is empty
                 final_draft = None
        else:
            print("⚠️ Workflow não disponível. Usando pipeline legado.")
            final_draft = None

        if final_draft:
            # Se o Workflow funcionou, retornamos direto (bypass legacy orchestration)
             return {
                "status": "success",
                "filename": filename,
                "analysis": final_draft,
                "model_used": f"Agentic Workflow V3 ({agent_provider})",
                "timestamp": time.time()
            }

        # 2. Run Pipeline (Legacy / Fallback)
        if mode == "v3" and keys:
            # V3: Autonomous Agent (Hybrid LangGraph Agents)
            if run_autonomous_magistrate is None:
                return {"error": "ERRO DE INSTALAÇÃO (V3): Engine Agente não disponível.", "filename": filename}
            
            # --- MIRROR STRATEGY FOR V3 (com Dossiê Forense) ---
            mirror_context = ""
            if template_files:
                 _dossier = generate_style_dossier(template_files, keys.get('google') or api_key)
                 mirror_context = retrieve_mirror_context(clean_content, keys.get('google') or api_key, template_files, style_dossier=_dossier)

            # Normalizar output para o formato esperado pelo front
            # returns (final_json, logs_list)
            v3_json, v3_logs = run_autonomous_magistrate(clean_content, keys)
            
            # Extract content safely
            final_minuta = v3_json.get("minuta_final", "Minuta não gerada.")
            reasoning = v3_json.get("fundamentacao_logica", "Raciocínio não disponível.")
            
            # Format reasoning string if it's a dict
            if isinstance(reasoning, dict):
                 reasoning = "\n".join([f"**{k}:** {v}" for k,v in reasoning.items()])
            
            results = {
                "final_report": final_minuta,
                "auditor_dashboard": "Auditoria Integrada ao Processo V3 (Ver Logs)",
                "style_report": "Gerado via Agentic Style Guide (V3)",
                "steps": {"logs": v3_logs},
                "diagnostic_reasoning": reasoning  # <--- NEW FIELD
            }
            
        elif mode == "v2" and keys:
            # V2: Ensemble Pipeline (Assembly Line)
            # Gemini -> DeepSeek -> Claude
            ensemble_output = run_ensemble_orchestration(clean_content, keys, template_files=template_files)
            results = ensemble_output # Já retorna no formato certo
            results["filename"] = filename # Garante filename
            
            # EXPOSE REASONING V2 (DeepSeek)
            results["diagnostic_reasoning"] = ensemble_output.get("steps", {}).get("analise_material", "Raciocínio não disponível.")
            
        else:
            # V1: Standard Pipeline (Flexible LLM)
            # V1: Standard Pipeline (Flexible LLM)
            # Precisamos mapear os parâmetros antigos para o novo formato de config
            # Se chamou via process_single_case_pipeline, api_key é a chave Google (default V1 legacy)
            # Para usar multi-model no batch, precisaremos passar 'keys' no futuro.
            # Por compatibilidade, se 'keys' não existir, assume Google Default.
            
            if keys:
                 # Novo formato: usa as chaves e modelos do keys se disponíveis
                 main_cfg = keys.get('v1_main_config', {'provider': 'google', 'model': 'gemini-3-pro-preview', 'key': api_key})
                 style_cfg = keys.get('v1_style_config', {'provider': 'google', 'model': 'gemini-3-flash-preview', 'key': api_key})
            else:
                 # Fallback legacy
                 main_cfg = {'provider': 'google', 'model': 'gemini-3-pro-preview', 'key': api_key}
                 style_cfg = {'provider': 'google', 'model': 'gemini-3-flash-preview', 'key': api_key}

            results = run_standard_orchestration(clean_content, main_cfg, style_cfg, status_callback=None, template_files=template_files, google_key=api_key, outline=outline, style_prompt=style_prompt)
        
        # === NORMALIZAÇÃO DO FINAL_REPORT ===
        # Garante que final_report seja sempre string (alguns modelos retornam lista)
        if results.get("final_report"):
            fr = results["final_report"]
            if isinstance(fr, list):
                # Extrai texto de estruturas como [{'type': 'text', 'text': '...'}]
                text_parts = []
                for item in fr:
                    if isinstance(item, dict) and 'text' in item:
                        text_parts.append(item['text'])
                    else:
                        text_parts.append(str(item))
                results["final_report"] = "\n".join(text_parts)
            elif not isinstance(fr, str):
                results["final_report"] = str(fr)
        
        # 3. Save Result
        report_id = hashlib.md5(f"{filename}_{time.time()}".encode()).hexdigest()
        
        # Add metadata for the UI
        results["filename"] = filename
        results["report_id"] = report_id
        results["timestamp"] = time.time()
        
        # Ensure directory exists
        os.makedirs("data/reports", exist_ok=True)
        
        with open(f"data/reports/{report_id}.json", "w") as f:
            json.dump(results, f, ensure_ascii=False)
            
        return {"report_id": report_id, "filename": filename, "status": "success"}

    except Exception as e:
        return {"error": str(e), "filename": filename}

def process_batch_parallel(files, api_key, template_files=None, text_cache_dict=None, progress_callback=None, mode="v1", keys=None, ocr_engine_choice="marker"):
    """
    Processa lista de arquivos EM SÉRIE (para evitar Rate Limit).
    Suporta V1/V2/V3 via worker.
    """
    # Pré-carrega bytes para evitar problemas de thread/cursor
    files_data = []
    for f in files:
        try:
            f.seek(0)
            c_text = None
            if text_cache_dict and f.name in text_cache_dict:
                c_text = text_cache_dict[f.name]
            
            # Leitura robusta de bytes
            content = f.read()
            if isinstance(content, str):
                content = content.encode('utf-8', errors='replace')
                
            files_data.append({
                "name": f.name,
                "bytes": content,
                "cached_text": c_text
            })
        except Exception as e:
            files_data.append({
                "name": f.name,
                "bytes": None,
                "cached_text": None,
                "error": f"Erro de Leitura (Upload): {str(e)}"
            })

    results_list = []
    total_files = len(files_data)
    
    print(f"🚀 Iniciando Batch Profundo (Série) para {total_files} arquivos...")
    os.makedirs("data/reports", exist_ok=True)

    # Execução em SÉRIE (Loop)
    # Motivo: Evitar Rate Limit do Google/OpenAI e uso excessivo de RAM com ChromaDB múltiplos
    for i, data in enumerate(files_data):
        try:
            filename = data["name"]
            print(f"🔄 Processando {i+1}/{total_files}: {filename}")
            
            if progress_callback:
                 progress_callback(i, total_files, filename) # 0-indexed start
            
            # Verifica se houve erro de leitura prévia
            if "error" in data:
                results_list.append({
                    "filename": filename,
                    "error": data["error"],
                    "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
                })
                continue
            
            # Chama o Pipeline Completo (OCR -> RAG -> Agentes)
            res = process_single_case_pipeline(
                pdf_bytes=data["bytes"], 
                filename=filename, 
                api_key=api_key, 
                template_files=template_files,
                cached_text=data["cached_text"],
                mode=mode,
                keys=keys,
                ocr_engine_choice=ocr_engine_choice
            )
            
            # Usa report_id gerado pelo process_single_case_pipeline (evita duplicação)
            res['mode'] = mode
            res['filename'] = filename
            if 'timestamp' not in res:
                res['timestamp'] = time.strftime("%Y-%m-%d %H:%M:%S")

            results_list.append(res)
            
            # Pausa entre processos para respeitar rate limit Azure S0 (tokens/min)
            time.sleep(10)

        except Exception as e:
            print(f"❌ Erro Crítico em {data['name']}: {e}")
            traceback.print_exc()
            results_list.append({
                "filename": data['name'],
                "error": str(e),
                "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
            })
            
    # Final Callback
    if progress_callback:
        progress_callback(total_files, total_files, "Concluído!")
                
    return results_list
