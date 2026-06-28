"""
flow_engine.py — Executes user-defined agent flows defined as JSON.

Suporta os seguintes tipos de nó:
  • start, end
  • agent  — chama LLM com prompt
  • router — bifurcação true/false via condição Python
  • switch — classificador multi-saída via LLM
  • hil    — Human-in-the-loop (pausa e emite evento, encerra stream)
  • docx   — converte markdown da entrada em arquivo .docx
  • juris  — RAG sobre acordãos do TJMG
  • modelo — RAG sobre templates do usuário
  • estilo — aplica style dossier do usuário

Substituições de variáveis: `{{Label do Nó}}` no prompt é substituído pela saída
do nó cujo label bate (case-insensitive).
"""

import ast
import json
import os
import re
import base64
import uuid
import tempfile
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Generator

import model_profiles as mp


_SAFE_BUILTINS = {
    "len": len, "int": int, "float": float, "str": str, "bool": bool,
    "abs": abs, "min": min, "max": max, "round": round,
}

# Router conditions come from a flow's JSON config, and flows are shareable and
# executed by *other* users. eval() with a restricted __builtins__ is NOT safe:
# ().__class__.__bases__[0].__subclasses__()[...] reaches arbitrary code without
# touching any builtin. So we gate the expression structurally before eval'ing.
_COND_ALLOWED_NODES = frozenset({
    ast.Expression, ast.BoolOp, ast.BinOp, ast.UnaryOp, ast.Compare, ast.IfExp,
    ast.Call, ast.keyword, ast.Name, ast.Load, ast.Constant,
    ast.List, ast.Tuple, ast.Dict, ast.Set, ast.Subscript, ast.Slice, ast.Starred,
    ast.Attribute,  # non-dunder only (enforced below)
    ast.ListComp, ast.SetComp, ast.DictComp, ast.GeneratorExp, ast.comprehension,
    ast.And, ast.Or, ast.Not, ast.Invert, ast.UAdd, ast.USub,
    ast.Add, ast.Sub, ast.Mult, ast.Div, ast.FloorDiv, ast.Mod, ast.Pow,
    ast.LShift, ast.RShift, ast.BitOr, ast.BitXor, ast.BitAnd, ast.MatMult,
    ast.Eq, ast.NotEq, ast.Lt, ast.LtE, ast.Gt, ast.GtE,
    ast.Is, ast.IsNot, ast.In, ast.NotIn,
})


def _condition_is_safe(condition: str):
    """Returns the parsed expression if structurally safe, else None.

    Rejects any node outside the allowlist and any '_'-prefixed identifier or
    attribute — which closes dunder-traversal sandbox escapes that a builtins
    restriction alone cannot.
    """
    try:
        tree = ast.parse(condition, mode="eval")
    except SyntaxError:
        return None
    for node in ast.walk(tree):
        if type(node) not in _COND_ALLOWED_NODES:
            return None
        if isinstance(node, ast.Attribute) and node.attr.startswith("_"):
            return None
        if isinstance(node, ast.Name) and node.id.startswith("_"):
            return None
    return tree


def _safe_eval(condition: str, state: dict) -> bool:
    tree = _condition_is_safe(condition)
    if tree is None:
        return False
    try:
        compiled = compile(tree, "<flow_condition>", "eval")
        return bool(eval(compiled, {"__builtins__": _SAFE_BUILTINS}, state))  # noqa: S307
    except Exception:
        return False


def _sse(data: dict) -> str:
    return f"data: {json.dumps(data, ensure_ascii=False)}\n\n"


def _build_label_map(nodes_list: list) -> dict[str, str]:
    """Retorna {label_normalizado: output_var} para substituição em prompts."""
    m: dict[str, str] = {}
    for n in nodes_list:
        cfg = n.get("data", n.get("config", {}))
        label = (cfg.get("label") or "").strip().lower()
        ovar = cfg.get("output_var") or n.get("id")
        if label and ovar:
            m[label] = ovar
    return m


def _substitute_vars(text: str, state: dict, label_map: dict[str, str]) -> str:
    """Substitui {{Label do Nó}} pelo conteúdo armazenado em state[output_var]."""
    if not text or "{{" not in text:
        return text

    def replace(match: re.Match) -> str:
        key = match.group(1).strip().lower()
        ovar = label_map.get(key)
        if ovar and ovar in state:
            return str(state[ovar])
        return match.group(0)

    return re.sub(r"\{\{\s*([^{}]+?)\s*\}\}", replace, text)


def _build_user_message(state: dict, input_text: str) -> str:
    """Constrói a mensagem do usuário com input principal + contexto de nós anteriores."""
    context = "\n\n".join(
        f"[{k}]:\n{v}" for k, v in state.items()
        if k != "input_text" and not k.startswith("__")
    )
    user_msg = f"Entrada principal:\n{state.get('input_text', input_text)}"
    if context:
        user_msg = f"{context}\n\n{user_msg}"
    return user_msg


# ─────────────────────────────────────────────────────────────────────
# NÓS ESPECIAIS — handlers individuais
# ─────────────────────────────────────────────────────────────────────

def _extract_usage(response) -> tuple[int, int]:
    """Lê tokens (input, output) da resposta LangChain (vários providers)."""
    try:
        meta = getattr(response, "usage_metadata", None)
        if isinstance(meta, dict):
            return int(meta.get("input_tokens", 0) or 0), int(meta.get("output_tokens", 0) or 0)
    except Exception:
        pass
    try:
        rm = getattr(response, "response_metadata", None) or {}
        usage = rm.get("token_usage") or rm.get("usage") or {}
        inp = usage.get("prompt_tokens") or usage.get("input_tokens") or 0
        out = usage.get("completion_tokens") or usage.get("output_tokens") or 0
        return int(inp or 0), int(out or 0)
    except Exception:
        return 0, 0


def _resolve_model(raw, default_role, state):
    """Resolve o modelo de um nó respeitando o perfil de custo ativo.

    - '@papel' (ex.: '@reasoner') → usa o perfil ativo em state['_profile'].
    - nome literal (ex.: 'gpt-5.3-chat') → usado direto (retrocompatível).
    - vazio/None → cai para o papel default informado, dentro do perfil ativo.
    """
    profile = state.get("_profile")
    if isinstance(raw, str) and raw.startswith("@"):
        return mp.resolve(raw[1:], profile)
    if raw:
        return raw
    return mp.resolve(default_role, profile)


def _run_agent(node, cfg, state, input_text, label_map):
    """Agente LLM padrão. Retorna (output_str, output_var, metrics)."""
    import backend as be
    from langchain_core.messages import SystemMessage, HumanMessage

    model = _resolve_model(cfg.get("model"), "default", state)
    prompt = cfg.get("prompt", "")
    knowledge = cfg.get("knowledge", "")
    output_var = cfg.get("output_var") or node["id"]

    # Substitui {{Label}} pelos valores reais
    prompt = _substitute_vars(prompt, state, label_map)

    llm = be.get_llm(model, temperature=0.3)

    full_prompt = prompt
    if knowledge:
        full_prompt = (
            (prompt + "\n\n" if prompt else "")
            + "BASE DE CONHECIMENTO (use como referência para responder):\n"
            + knowledge[:80_000]
        )

    msgs = []
    if full_prompt:
        msgs.append(SystemMessage(content=full_prompt))
    msgs.append(HumanMessage(content=_build_user_message(state, input_text)))

    response = llm.invoke(msgs)
    result = response.content if hasattr(response, "content") else str(response)
    inp_tok, out_tok = _extract_usage(response)
    return result, output_var, {"model": model, "input_tokens": inp_tok, "output_tokens": out_tok}


def _run_switch(node, cfg, state, input_text, label_map):
    """Classificador LLM multi-saída. Retorna a categoria escolhida."""
    import backend as be
    from langchain_core.messages import SystemMessage, HumanMessage

    model = _resolve_model(cfg.get("model"), "classifier", state)
    raw_cats = cfg.get("categories", "")
    categories = [c.strip() for c in raw_cats.split("|") if c.strip()]
    if not categories:
        raise ValueError("Switch sem categorias definidas.")

    llm = be.get_llm(model, temperature=0.0)
    cats_str = "\n".join(f"- {c}" for c in categories)
    sys_msg = SystemMessage(content=(
        "Você é um classificador. Analise a entrada e responda APENAS com "
        "exatamente um item da lista de categorias abaixo, sem nenhum texto extra:\n\n"
        f"{cats_str}"
    ))
    user_msg = HumanMessage(content=_build_user_message(state, input_text))
    response = llm.invoke([sys_msg, user_msg])
    raw = (response.content if hasattr(response, "content") else str(response)).strip()

    # Encontra a categoria que melhor bate (case-insensitive, prefix)
    chosen = next((c for c in categories if c.lower() in raw.lower()), categories[0])
    inp_tok, out_tok = _extract_usage(response)
    return chosen, {"model": model, "input_tokens": inp_tok, "output_tokens": out_tok}


def _run_docx(node, cfg, state, input_text, label_map):
    """Converte o markdown da última saída em .docx, salva em /tmp e devolve URL."""
    from docx import Document  # python-docx
    from docx.shared import Pt

    # Pega o último conteúdo gerado (output do nó anterior)
    last_output = ""
    for k, v in list(state.items())[::-1]:
        if k != "input_text" and not k.startswith("__"):
            last_output = str(v)
            break
    if not last_output:
        last_output = state.get("input_text", input_text)

    filename = cfg.get("filename") or "documento.docx"
    if not filename.endswith(".docx"):
        filename += ".docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for raw_line in last_output.split("\n"):
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            p = doc.add_paragraph()
            # negrito simples **texto**
            parts = re.split(r"(\*\*[^*]+\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    p.add_run(part)

    out_dir = os.path.join(tempfile.gettempdir(), "jurisbusca_docx")
    os.makedirs(out_dir, exist_ok=True)
    file_id = uuid.uuid4().hex[:12]
    out_path = os.path.join(out_dir, f"{file_id}__{filename}")
    doc.save(out_path)

    url = f"/api/flows/docx-download/{file_id}__{filename}"
    return f"📄 **Documento gerado:** [{filename}]({url})", "docx_url", url


def _run_juris(node, cfg, state, input_text, label_map):
    """Busca semântica em acordãos TJMG via jurisprudence_search."""
    query = (cfg.get("query") or "").strip()
    query = _substitute_vars(query, state, label_map)
    if not query:
        query = state.get("input_text", input_text)[:500]
    top_k = int(cfg.get("top_k") or 5)

    results = []
    try:
        import jurisprudence_search as js
        if not js.is_available():
            return "(banco de jurisprudência indisponível)", cfg.get("output_var") or node["id"]
        # tenta semantic_search; cai pra search FTS se embeddings indisponíveis
        try:
            hits = js.semantic_search(query, limit=top_k)
        except Exception:
            hits = js.search(query, limit=top_k)
        for i, h in enumerate(hits or []):
            num = h.get("numero_processo") or h.get("id") or "—"
            tipo = h.get("tipo_recurso") or "—"
            ementa = h.get("ementa") or h.get("text") or ""
            results.append(f"### Acordão {i+1} — {tipo} {num}\n{str(ementa)[:1200]}")
    except Exception as exc:
        results = [f"(erro ao buscar jurisprudência: {exc})"]

    text = "\n\n".join(results) if results else "(nenhum resultado)"
    return text, cfg.get("output_var") or node["id"]


def _run_modelo(node, cfg, state, input_text, label_map):
    """Busca semântica nos templates do usuário via backend.search_templates."""
    import backend as be

    query = (cfg.get("query") or "").strip()
    query = _substitute_vars(query, state, label_map)
    if not query:
        query = state.get("input_text", input_text)[:500]
    top_k = int(cfg.get("top_k") or 3)
    user_id = state.get("__user_id", "default")
    token = state.get("__token", "")

    try:
        hits = be.search_templates(user_id, query, k=top_k, token=token)
        if not hits:
            return "(nenhum modelo encontrado)", cfg.get("output_var") or node["id"]
        text = "\n\n---\n\n".join(
            f"### Modelo {i+1}: {h.get('filename') or h.get('name') or 'sem nome'}\n{(h.get('text') or h.get('content') or '')[:1500]}"
            for i, h in enumerate(hits)
        )
    except Exception as exc:
        text = f"(erro ao buscar modelos: {exc})"

    return text, cfg.get("output_var") or node["id"]


def _run_extractor(node, cfg, state, input_text, label_map):
    """Extrai dados estruturados em JSON. Cada chave vira variável no state.

    Config esperado:
      - fields: string "nome:tipo:descricao|nome:tipo:descricao"
        tipos suportados: string, number, integer, boolean, array, object
      - model: opcional, default gpt-5.4-mini
    """
    import backend as be
    from langchain_core.messages import SystemMessage, HumanMessage

    fields_raw = (cfg.get("fields") or "").strip()
    if not fields_raw:
        raise ValueError("Extrator sem campos definidos.")

    # Parse fields
    fields: list[dict] = []
    for line in fields_raw.split("|"):
        line = line.strip()
        if not line:
            continue
        parts = [p.strip() for p in line.split(":", 2)]
        name = parts[0]
        ftype = parts[1] if len(parts) > 1 else "string"
        desc = parts[2] if len(parts) > 2 else ""
        fields.append({"name": name, "type": ftype.lower(), "desc": desc})

    if not fields:
        raise ValueError("Nenhum campo válido configurado.")

    # Constrói descrição do schema para o prompt
    schema_lines = []
    for f in fields:
        line = f'- "{f["name"]}" ({f["type"]})'
        if f["desc"]:
            line += f": {f['desc']}"
        schema_lines.append(line)
    schema_desc = "\n".join(schema_lines)

    # Build user_msg
    user_input = _build_user_message(state, input_text)

    model = _resolve_model(cfg.get("model"), "classifier", state)
    llm = be.get_llm(model, temperature=0.0)

    sys_prompt = (
        "Você é um extrator de dados estruturados. Analise a entrada abaixo e devolva "
        "EXCLUSIVAMENTE um objeto JSON válido (sem comentários, sem texto antes/depois, "
        "sem markdown fence) com os seguintes campos:\n\n"
        f"{schema_desc}\n\n"
        "Regras:\n"
        "- Use null quando o campo não puder ser determinado a partir da entrada.\n"
        "- Para arrays/objects, devolva sintaxe JSON correta.\n"
        "- Respeite estritamente os nomes dos campos (case-sensitive)."
    )

    # tentativa + 2 retries com feedback
    last_err = ""
    parsed: dict | None = None
    total_in = 0
    total_out = 0
    for attempt in range(3):
        msgs = [SystemMessage(content=sys_prompt), HumanMessage(content=user_input)]
        if attempt > 0 and last_err:
            msgs.append(HumanMessage(content=(
                f"A tentativa anterior falhou na validação: {last_err}\n"
                "Devolva agora apenas o JSON corrigido, sem texto extra."
            )))
        response = llm.invoke(msgs)
        inp_t, out_t = _extract_usage(response)
        total_in += inp_t
        total_out += out_t
        raw = (response.content if hasattr(response, "content") else str(response)).strip()

        # remove possível ```json ... ``` wrap
        raw_clean = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw, flags=re.MULTILINE).strip()

        try:
            parsed = json.loads(raw_clean)
            # validação leve por tipo
            for f in fields:
                if f["name"] not in parsed:
                    raise ValueError(f"campo obrigatório '{f['name']}' ausente")
                val = parsed[f["name"]]
                if val is None:
                    continue
                t = f["type"]
                if t == "string" and not isinstance(val, str):
                    raise ValueError(f"'{f['name']}' deveria ser string, recebido {type(val).__name__}")
                if t in ("number", "integer") and not isinstance(val, (int, float)):
                    raise ValueError(f"'{f['name']}' deveria ser número")
                if t == "boolean" and not isinstance(val, bool):
                    raise ValueError(f"'{f['name']}' deveria ser boolean")
                if t == "array" and not isinstance(val, list):
                    raise ValueError(f"'{f['name']}' deveria ser array")
                if t == "object" and not isinstance(val, dict):
                    raise ValueError(f"'{f['name']}' deveria ser object")
            break  # sucesso
        except Exception as exc:
            last_err = str(exc)
            parsed = None
            continue

    if parsed is None:
        raise ValueError(f"Falha em extrair JSON válido após 3 tentativas. Último erro: {last_err}")

    # Mergeo cada chave como variável separada no state
    # (acessível por @nome_campo nos próximos nós)
    extracted_state_updates: dict = {}
    for f in fields:
        v = parsed.get(f["name"])
        if v is None:
            extracted_state_updates[f["name"]] = ""
        elif isinstance(v, (list, dict)):
            extracted_state_updates[f["name"]] = json.dumps(v, ensure_ascii=False)
        else:
            extracted_state_updates[f["name"]] = v

    # devolve via output_var um JSON formatado para exibição;
    # também marca a variável bruta com sufixo _json
    output_var = cfg.get("output_var") or node["id"]
    formatted_json = json.dumps(parsed, ensure_ascii=False, indent=2)

    # injeta as variáveis no state DIRETAMENTE — o caller precisa atualizar isso
    # mas como _execute_simple_node retorna um dict de updates, vou retornar
    # tudo agregado: as keys individuais + a versão JSON formatada
    extracted_state_updates[output_var] = formatted_json
    extracted_state_updates[f"{output_var}_json"] = formatted_json
    extracted_state_updates["__metrics__"] = {"model": model, "input_tokens": total_in, "output_tokens": total_out}

    return formatted_json, output_var, extracted_state_updates


def _run_subflow(node, cfg, state, input_text, label_map):
    """Carrega outro fluxo salvo e executa recursivamente.

    Limita profundidade via state["__subflow_depth"] (max 4 níveis).
    """
    sub_flow_id = (cfg.get("flow_id") or "").strip()
    if not sub_flow_id:
        raise ValueError("Sub-fluxo sem flow_id configurado.")

    depth = int(state.get("__subflow_depth", 0))
    if depth >= 4:
        raise ValueError("Profundidade máxima de sub-fluxos excedida (possível loop).")

    from history_db import get_flow, get_flow_shared
    user_id = state.get("__user_id", "")
    email = state.get("__email", "")
    sub_flow = get_flow(sub_flow_id, user_id) or get_flow_shared(sub_flow_id, email)
    if not sub_flow:
        raise ValueError(f"Sub-fluxo {sub_flow_id} não encontrado ou sem acesso.")

    # Input do sub-fluxo: última saída do estado atual, senão input_text
    sub_input = ""
    for k, v in list(state.items())[::-1]:
        if k != "input_text" and not k.startswith("__"):
            sub_input = str(v)
            break
    if not sub_input:
        sub_input = state.get("input_text", input_text)

    # Estado inicial do sub-fluxo: meta keys + nada mais (isolamento)
    sub_state = {k: v for k, v in state.items() if k.startswith("__")}
    sub_state["__subflow_depth"] = depth + 1

    # Executa silenciosamente, captura saída final
    final_output = ""
    nodes_done = 0
    for ev_str in build_and_run_flow(sub_flow["config"], sub_input, initial_state=sub_state):
        if ev_str.startswith("data: "):
            try:
                ev = json.loads(ev_str[6:].strip())
                if ev.get("event") == "node_done":
                    nodes_done += 1
                if ev.get("event") == "flow_done":
                    final_output = ev.get("output", "")
            except Exception:
                pass

    label = cfg.get("label") or sub_flow.get("name") or "Sub-fluxo"
    output = f"_{nodes_done} etapas executadas em **{sub_flow.get('name', label)}**_\n\n{final_output}"
    return output, cfg.get("output_var") or node["id"]


def _run_estilo(node, cfg, state, input_text, label_map):
    """Reescreve o último texto aplicando o style dossier do usuário."""
    import backend as be
    from langchain_core.messages import SystemMessage, HumanMessage

    style_dossier = state.get("__style_dossier", "")
    last_output = ""
    for k, v in list(state.items())[::-1]:
        if k != "input_text" and not k.startswith("__"):
            last_output = str(v)
            break
    if not last_output:
        last_output = state.get("input_text", input_text)

    if not style_dossier:
        return last_output + "\n\n*(style dossier não configurado — texto retornado sem ajuste)*", cfg.get("output_var") or node["id"]

    llm = be.get_llm(_resolve_model(cfg.get("model"), "style", state), temperature=0.4)
    sys_msg = SystemMessage(content=(
        "Você é um editor de texto jurídico. Reescreva o texto abaixo aplicando "
        "rigorosamente o estilo de escrita descrito a seguir, sem alterar o conteúdo "
        "factual ou os dispositivos legais citados.\n\n"
        f"ESTILO ALVO:\n{style_dossier[:20_000]}"
    ))
    response = llm.invoke([sys_msg, HumanMessage(content=last_output)])
    result = response.content if hasattr(response, "content") else str(response)
    return result, cfg.get("output_var") or node["id"]


# ─────────────────────────────────────────────────────────────────────
# Worker thread-safe para nós "simples" (executáveis em paralelo)
# ─────────────────────────────────────────────────────────────────────

SIMPLE_NODE_TYPES = {"agent", "juris", "modelo", "estilo", "docx", "extractor", "subflow"}

_SPECIAL_HANDLERS = {
    "juris": _run_juris,
    "modelo": _run_modelo,
    "estilo": _run_estilo,
    "subflow": _run_subflow,
}


def _execute_simple_node(node, state_snapshot, input_text, label_map):
    """Executa um nó simples e retorna (eventos, state_updates).

    Roda em qualquer thread. Recebe um snapshot do state (cópia) e devolve só
    as variáveis novas que devem ser mergidas no state global.
    """
    import time as _t
    import flow_pricing as _fp

    node_id = node["id"]
    ntype = node["type"]
    cfg = node.get("data", node.get("config", {}))
    label = cfg.get("label") or node_id

    events: list[dict] = [{"event": "node_start", "node_id": node_id, "label": label}]
    updates: dict = {}
    metrics: dict = {"model": "", "input_tokens": 0, "output_tokens": 0, "cost_usd": 0.0}
    t0 = _t.perf_counter()

    try:
        if ntype == "agent":
            output, ovar, m = _run_agent(node, cfg, state_snapshot, input_text, label_map)
            metrics["model"] = m.get("model", "")
            metrics["input_tokens"] = m.get("input_tokens", 0)
            metrics["output_tokens"] = m.get("output_tokens", 0)
            metrics["cost_usd"] = _fp.estimate_cost_usd(
                metrics["model"], metrics["input_tokens"], metrics["output_tokens"]
            )
            updates[ovar] = output
            events.append({"event": "node_done", "node_id": node_id, "label": label,
                           "output": output, "output_var": ovar, **metrics})

        elif ntype == "docx":
            output, ovar, url = _run_docx(node, cfg, state_snapshot, input_text, label_map)
            updates[ovar] = url
            updates["__docx_url"] = url
            events.append({"event": "node_done", "node_id": node_id, "label": label,
                           "output": output, "output_var": ovar, "download_url": url, **metrics})

        elif ntype == "extractor":
            output, ovar, extra_updates = _run_extractor(node, cfg, state_snapshot, input_text, label_map)
            m = extra_updates.pop("__metrics__", None) if isinstance(extra_updates, dict) else None
            updates.update(extra_updates)  # cada chave do JSON vira variável
            if m:
                metrics["model"] = m.get("model", "")
                metrics["input_tokens"] = m.get("input_tokens", 0)
                metrics["output_tokens"] = m.get("output_tokens", 0)
                metrics["cost_usd"] = _fp.estimate_cost_usd(
                    metrics["model"], metrics["input_tokens"], metrics["output_tokens"]
                )
            events.append({"event": "node_done", "node_id": node_id, "label": label,
                           "output": output, "output_var": ovar,
                           "extracted_fields": list(extra_updates.keys()), **metrics})

        elif ntype in _SPECIAL_HANDLERS:
            output, ovar = _SPECIAL_HANDLERS[ntype](node, cfg, state_snapshot, input_text, label_map)
            updates[ovar] = output
            events.append({"event": "node_done", "node_id": node_id, "label": label,
                           "output": output, "output_var": ovar, **metrics})

    except Exception as exc:
        duration_ms = int((_t.perf_counter() - t0) * 1000)
        events.append({"event": "node_error", "node_id": node_id, "label": label,
                       "error": str(exc), "duration_ms": duration_ms, **metrics})
        updates["__error__"] = str(exc)
        return events, updates

    # Adiciona duration_ms ao evento node_done
    duration_ms = int((_t.perf_counter() - t0) * 1000)
    if events and events[-1].get("event") == "node_done":
        events[-1]["duration_ms"] = duration_ms

    return events, updates


# ─────────────────────────────────────────────────────────────────────
# MAIN: build_and_run_flow
# ─────────────────────────────────────────────────────────────────────

def build_and_run_flow(
    flow_config: dict,
    input_text: str,
    *,
    start_from: str | None = None,
    initial_state: dict | None = None,
) -> Generator[str, None, None]:
    """Executa o fluxo com scheduler topológico (suporta fan-out paralelo)."""
    nodes_list = flow_config.get("nodes", [])
    nodes = {n["id"]: n for n in nodes_list}
    edges = flow_config.get("edges", [])
    label_map = _build_label_map(nodes_list)

    # adjacency: source -> list of (target, label, source_handle)
    adj: dict[str, list[tuple[str, str, str]]] = {}
    # predecessors: target -> list of (source, source_handle)
    predecessors: dict[str, list[tuple[str, str]]] = {n["id"]: [] for n in nodes_list}
    for e in edges:
        src, tgt = e["source"], e["target"]
        h = e.get("sourceHandle", "") or ""
        adj.setdefault(src, []).append((tgt, e.get("label", ""), h))
        predecessors[tgt].append((src, h))

    state: dict = dict(initial_state or {})
    state.setdefault("input_text", input_text)
    state_lock = threading.Lock()

    completed: set[str] = set()
    cancelled: set[str] = set()

    # Pre-marca como completed: tudo na initial_state (resume) que já tem chave igual a output_var
    if start_from and initial_state:
        for n in nodes_list:
            ovar = (n.get("data") or {}).get("output_var") or n["id"]
            if ovar in state and n["id"] != start_from:
                completed.add(n["id"])

    # determina start
    start_node = next((n for n in nodes_list if n["type"] == "start"), None)
    start_id = start_from or (start_node["id"] if start_node else None)
    if not start_id:
        yield _sse({"event": "error", "message": "Nenhum nó de início encontrado no fluxo."})
        return

    # Para resume, marca todos os ancestrais de start_from como completed.
    # Sem isso, nós sem output_var em state (ex.: hil) seriam re-executados
    # pelo scheduler e o fluxo entraria em loop emitindo human_required.
    if start_from:
        if start_node:
            completed.add(start_node["id"])
        stack = [start_from]
        visited: set[str] = set()
        while stack:
            n_id = stack.pop()
            if n_id in visited:
                continue
            visited.add(n_id)
            for pred_id, _h in predecessors.get(n_id, []):
                if pred_id != start_from:
                    completed.add(pred_id)
                stack.append(pred_id)

    def _cancel_subtree(skipped_targets: list[str]):
        """Propaga cancelamento: nós cujos predecessors todos viraram cancelled."""
        stack = list(skipped_targets)
        while stack:
            n = stack.pop()
            if n in cancelled or n in completed:
                continue
            preds = predecessors.get(n, [])
            if preds and all(p in cancelled or (p, _) in [] for p, _ in preds):
                # se todos os predecessors estão cancelled, cancela este também
                if all(p in cancelled for p, _ in preds):
                    cancelled.add(n)
                    for tgt, _l, _h in adj.get(n, []):
                        stack.append(tgt)

    def _ready_nodes() -> list[str]:
        out: list[str] = []
        for n_id in nodes:
            if n_id in completed or n_id in cancelled:
                continue
            preds = predecessors[n_id]
            if not preds:
                if n_id == start_id and start_id not in completed:
                    out.append(n_id)
                continue
            all_decided = all(p in completed or p in cancelled for p, _ in preds)
            any_alive = any(p in completed for p, _ in preds)
            if all_decided and any_alive:
                out.append(n_id)
        return out

    def _emit_final():
        # Preferência: usar o output_var do nó imediatamente anterior ao "end"
        # (a verdadeira "saída final" do fluxo, respeitando o ramo escolhido).
        final_output: object | None = None
        end_node_local = next((n for n in nodes_list if n["type"] == "end"), None)
        if end_node_local is not None:
            preds = predecessors.get(end_node_local["id"], [])
            # entre os predecessors, prioriza os que realmente rodaram (completed)
            ordered_preds = sorted(
                preds, key=lambda p: 0 if p[0] in completed else 1
            )
            for pred_id, _h in ordered_preds:
                if pred_id in cancelled:
                    continue
                pred_node = nodes.get(pred_id) or {}
                pred_cfg = pred_node.get("data", pred_node.get("config", {}))
                ovar = pred_cfg.get("output_var") or pred_id
                if ovar in state:
                    final_output = state[ovar]
                    break
        # Fallback: último valor não-meta gravado no state
        if final_output is None:
            final_output = state.get("input_text", input_text)
            for k, v in state.items():
                if k != "input_text" and not k.startswith("__"):
                    final_output = v
        clean_state = {k: v for k, v in state.items() if not k.startswith("__")}
        return _sse({"event": "flow_done", "output": final_output, "state": clean_state})

    # ── Loop principal ──────────────────────────────────────────────
    safety = 0
    while True:
        safety += 1
        if safety > 200:  # circuit breaker
            yield _sse({"event": "error", "message": "Loop overflow no fluxo."})
            return

        ready = _ready_nodes()
        if not ready:
            break

        # processa nós "control flow" um a um (não dá pra paralelizar decisões)
        control_ids = [r for r in ready if nodes[r]["type"] in ("start", "end", "router", "switch", "hil")]
        simple_ids = [r for r in ready if nodes[r]["type"] in SIMPLE_NODE_TYPES]

        # ── 1. START / END / ROUTER / SWITCH / HIL — sequencial ────
        if control_ids:
            current_id = control_ids[0]
            node = nodes[current_id]
            ntype = node["type"]
            cfg = node.get("data", node.get("config", {}))
            label = cfg.get("label") or current_id

            if ntype == "start":
                completed.add(current_id)
                continue

            if ntype == "end":
                yield _emit_final()
                return

            if ntype == "router":
                condition = cfg.get("condition", "True")
                yield _sse({"event": "node_start", "node_id": current_id, "label": label})
                branch = "true" if _safe_eval(condition, state) else "false"
                yield _sse({"event": "node_done", "node_id": current_id, "label": label, "branch": branch})
                targets = adj.get(current_id, [])
                # Matching robusto: sourceHandle ("true"/"false") OU label em pt-br
                # ("verdadeiro"/"falso") OU label igual ao branch.
                pt_alias = {"true": ("verdadeiro", "verdadeiro"), "false": ("falso", "falso")}
                pt_label = pt_alias[branch][0]
                chosen_target = next(
                    (tgt for tgt, lbl, h in targets
                     if (h or "").lower() == branch
                     or (lbl or "").strip().lower() == pt_label
                     or (lbl or "").strip().lower() == branch),
                    None,
                )
                if chosen_target is None:
                    if len(targets) == 1:
                        chosen_target = targets[0][0]
                    elif targets:
                        # Ramo "true"/"false" não está conectado — fluxo travaria
                        # silenciosamente; emite erro para o usuário entender.
                        yield _sse({
                            "event": "node_error", "node_id": current_id, "label": label,
                            "error": (
                                f"Roteador sem conexão para o ramo '{branch}'. "
                                "Conecte as saídas Verde (verdadeiro) e Laranja (falso) do nó."
                            ),
                        })
                        return
                # cancela os outros
                for tgt, _l, _h in targets:
                    if tgt != chosen_target:
                        cancelled.add(tgt)
                        _cancel_subtree([tgt])
                completed.add(current_id)
                continue

            if ntype == "switch":
                yield _sse({"event": "node_start", "node_id": current_id, "label": label})
                import time as _t
                t0 = _t.perf_counter()
                try:
                    chosen, switch_metrics = _run_switch(node, cfg, state, input_text, label_map)
                    state[cfg.get("output_var") or current_id] = chosen
                except Exception as exc:
                    duration_ms = int((_t.perf_counter() - t0) * 1000)
                    yield _sse({"event": "node_error", "node_id": current_id, "label": label,
                                "error": str(exc), "duration_ms": duration_ms})
                    return
                duration_ms = int((_t.perf_counter() - t0) * 1000)
                import flow_pricing as _fp
                cost_usd = _fp.estimate_cost_usd(
                    switch_metrics["model"], switch_metrics["input_tokens"], switch_metrics["output_tokens"]
                )
                yield _sse({"event": "node_done", "node_id": current_id, "label": label,
                            "branch": chosen, "output": chosen,
                            "duration_ms": duration_ms,
                            "model": switch_metrics["model"],
                            "input_tokens": switch_metrics["input_tokens"],
                            "output_tokens": switch_metrics["output_tokens"],
                            "cost_usd": cost_usd})
                targets = adj.get(current_id, [])
                chosen_target = next(
                    (tgt for tgt, _l, h in targets if h.lower() == chosen.lower()),
                    targets[0][0] if targets else None,
                )
                for tgt, _l, _h in targets:
                    if tgt != chosen_target:
                        cancelled.add(tgt)
                        _cancel_subtree([tgt])
                completed.add(current_id)
                continue

            if ntype == "hil":
                yield _sse({"event": "node_start", "node_id": current_id, "label": label})
                last_output = state.get("input_text", input_text)
                for k, v in state.items():
                    if k != "input_text" and not k.startswith("__"):
                        last_output = v
                targets = adj.get(current_id, [])
                next_id = targets[0][0] if targets else None
                yield _sse({
                    "event": "human_required",
                    "node_id": current_id,
                    "label": label,
                    "question": cfg.get("question") or "Aprove para continuar.",
                    "content": last_output,
                    "next_node_id": next_id,
                    "state": {k: v for k, v in state.items() if not k.startswith("__")},
                })
                return

        # ── 2. nós simples — paralelo se mais de 1 ─────────────────
        elif simple_ids:
            # snapshot atual do state (cada thread recebe sua cópia)
            with state_lock:
                snapshot = dict(state)

            if len(simple_ids) == 1:
                # single — roda inline (sem overhead de thread)
                events, updates = _execute_simple_node(
                    nodes[simple_ids[0]], snapshot, input_text, label_map
                )
                for ev in events:
                    yield _sse(ev)
                if "__error__" in updates:
                    return
                with state_lock:
                    state.update(updates)
                completed.add(simple_ids[0])
            else:
                # FAN-OUT paralelo
                yield _sse({
                    "event": "parallel_start",
                    "node_ids": simple_ids,
                    "labels": [nodes[n].get("data", {}).get("label") or n for n in simple_ids],
                })
                with ThreadPoolExecutor(max_workers=min(len(simple_ids), 8)) as pool:
                    futures = {
                        pool.submit(_execute_simple_node, nodes[n_id], snapshot, input_text, label_map): n_id
                        for n_id in simple_ids
                    }
                    for fut in as_completed(futures):
                        n_id = futures[fut]
                        events, updates = fut.result()
                        for ev in events:
                            yield _sse(ev)
                        if "__error__" in updates:
                            return
                        with state_lock:
                            state.update(updates)
                        completed.add(n_id)
                yield _sse({"event": "parallel_done"})
        else:
            # ready vazio depois das filtragens — algo errado
            break

    yield _emit_final()
